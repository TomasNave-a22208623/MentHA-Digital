
from django.shortcuts import render, redirect
from django.forms import ModelForm, TextInput, Textarea
from .models import *
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from .forms import *
from datetime import datetime

from .functions import *
from .decorators import *

from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.contrib.auth.models import Group as DjangoGroup

# Create your views here.

from django.template.defaulttags import register

import qrcode
import qrcode.image.svg
from io import BytesIO

# imports para imprimir grafico
from matplotlib import pyplot as plt
import io
import urllib, base64
import matplotlib

from django.shortcuts import render

from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer

from docx import Document
import win32com.client as win32
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pythoncom
import tempfile
from django.core.files import File

matplotlib.use('Agg')


# Para permitir acesso a views por grupo
# @user_passes_test(lambda u: u.groups.filter(name='YourGroupName').exists())

def streams(request):
    return render(request,'diario/streams.html')

@register.filter
def get_item(dictionary, key):
    return dictionary.get(key)


def nextSession(request):
    datas = SessaoDoGrupo.objects.exclude(data=None)
    datas = datas.filter(estado='PR')

    if bool(datas) == True:
        datas = datas.filter(estado='PR').order_by('data')[0]

    factory = qrcode.image.svg.SvgImage
    uri = request.build_absolute_uri('zoom')
    uri = uri.replace('abrirZ', 'z')
    img = qrcode.make(uri, image_factory=factory, box_size=5)
    img_pop = qrcode.make(uri, image_factory=factory, box_size=45)
    stream = BytesIO()
    stream_pop = BytesIO()
    img.save(stream)
    img_pop.save(stream_pop)

    contexto = {
        'proxima': datas,
        'ss': bool(datas),
        'svg': stream.getvalue().decode(),
        'svg_pop': stream_pop.getvalue().decode(),
    }

    return render(request, 'diario/nextSession.html', contexto)

def get_grupos(user):
    dinamizador = DinamizadorConvidado.objects.filter(user=user).first()
    mentor = Mentor.objects.filter(user=user).first()
    participante = Participante.objects.filter(user=user).first()
    sg = None
    is_participante = False
    if dinamizador:
        grupos = dinamizador.grupo.all()
    if mentor:
        grupos = mentor.grupo.all()
    if participante:
        grupos = participante.grupo.all()
        sg = SessaoDoGrupo.objects.filter(grupo__in=participante.grupo.all()).exclude(parte_ativa__isnull=True)
        is_participante = True
    if user.is_superuser:
        grupos = Grupo.objects.all()

        return grupos, sg, is_participante

def get_proxima_sessao(grupos):
    datas = SessaoDoGrupo.objects.exclude(data=None)
    #Caso um dinamizador/mentor tenha mais do que um grupo
    datas = datas.filter(estado__in=['PR','EC'], grupo__in=grupos)
    
    tem_proxima = False
    if len(datas) > 0:
        datas = datas.order_by('data')[0]
        tem_proxima = True
    
    return tem_proxima, datas
    
@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def dashboard(request):
    # doctor = request.user
    formGrupo = GrupoForm(request.POST or None)
    if formGrupo.is_valid():
        formGrupo.save()
        return redirect('diario:new_group')

    # Retrieve all objects without a name
    cuidadores_without_name = Cuidador.objects.filter(info_sensivel__nome__isnull=True)

    # Delete the objects without a name
    cuidadores_without_name.delete()
    cuidador = Cuidador.objects.filter(user=request.user).first()
    
    if cuidador:
        return redirect('diario:user_dashboard')
    
    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)

    factory = qrcode.image.svg.SvgImage
    uri = request.build_absolute_uri('zoom')
    uri = uri.replace('abrirZ', 'z')
    img = qrcode.make(uri, image_factory=factory, box_size=5)
    img_pop = qrcode.make(uri, image_factory=factory, box_size=45)
    stream = BytesIO()
    stream_pop = BytesIO()
    img.save(stream)
    img_pop.save(stream_pop)
    print(datas)
    contexto = {
    # 'grupos': Grupo.objects.filter(doctor=doctor),
    # Apagar a linha de baixo ao descomentar a linha de cima
    # next(obj for obj in DinamizadorConvidado.objects.all() if obj.user.username == request.user.username).grupo.all()
    'grupos': grupos,
    'cuidadores': Cuidador.objects.filter(grupo=None),
    'formGrupo': formGrupo,
    'proxima': datas,
    'tem_proxima': tem_proxima,
    'svg': stream.getvalue().decode(),
    'svg_pop': stream_pop.getvalue().decode(),
    }

    if is_participante and len(sg) > 1:
        participante = Participante.objects.filter(user=request.user).first()
        sg = sg.get()
        parte = sg.parte_ativa
        contexto['parte'] = parte
        contexto['sg'] = sg

        form_list = []
        lista_ids_escolhas_multiplas = []
        for pergunta in parte.perguntas.all():
            initial_data = {}
            r = Resposta.objects.filter(
                participante__id=participante.id,
                sessao_grupo__id=sg.id,
                pergunta=pergunta,
                parte_exercicio=parte,
            )

            if len(r) > 0:
                r = r.get()
                initial_data = {
                    'resposta_escrita': r.resposta_escrita,
                    'certo': r.certo,
                }
                if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                    lista_ids_escolhas_multiplas.append(r.resposta_escolha.id)

            if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                form = RespostaForm_RespostaEscrita(None, initial=initial_data)
            elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                form = RespostaForm_RespostaSubmetida(None)
            if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                form = None

            tuplo = (pergunta, parte, form)
            form_list.append(tuplo)
            contexto['form_list'] = form_list
            contexto['lista_ids_escolhas_multiplas'] = lista_ids_escolhas_multiplas
        return render(request, 'diario/parte_ativa.html', contexto)
    return render(request, 'diario/dashboard.html', contexto)


# @login_required(login_url='diario:login')
# @check_user_able_to_see_page('Todos')
# def menu_view(request):
#   return render(request,'mentha/app-list.html')


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def parte_ativa(request, sg_id):
    sg = SessaoDoGrupo.objects.get(id=sg_id)
    contexto = {}
    contexto['parte'] = sg.parte_ativa
    contexto['sg'] = sg

    respostas_existentes = {}
    lista_ids_escolhas_multiplas = []

    parte_ativa = sg.parte_ativa

    if not parte_ativa:
        sem_parte_ativa = True
        return render(request, 'diario/parte_ativa.html', contexto)

    parte = Parte_Exercicio.objects.get(id=parte_ativa.id)

    form_list = []  
    initial_data = {}
    
    if parte.perguntas.all():
        for pergunta in parte.perguntas.all():
            r = Resposta.objects.filter(
                participante__user=request.user,
                sessao_grupo=sg,
                pergunta=pergunta,
                # parte_grupo__id = parte_do_grupo_id,
                parte_exercicio=parte,
            )

            if len(r) > 0:
                r = r.get()
                initial_data = {
                    'resposta_escrita': r.resposta_escrita
                }
                if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                    lista_ids_escolhas_multiplas.append(r.resposta_escolha.id)

            if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                form = RespostaForm_RespostaEscrita(None, initial=initial_data)
            elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                form = RespostaForm_RespostaSubmetida(None)
            elif pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                form = None

            tuplo = (pergunta, parte.ordem, form)
            form_list.append(tuplo)

    contexto['form_list'] = form_list

    return render(request, 'diario/parte_ativa.html', contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def new_group(request):
    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)

    formGrupo = GrupoForm(request.POST or None)
    if formGrupo.is_valid():
        formGrupo.save()
        return HttpResponseRedirect(reverse('diario:dashboard_Care'))

    # Obter campos para filtar por (CARE)
    cuidadores = Cuidador.objects.all()
    filtrados_care = cuidadores.filter(grupo=None)

    conjunto_doencas = set()
    for cuidador in cuidadores:
        conjunto_doencas.update(cuidador.doencas_object)


    lista_pesquisa_cuidadores = {
        'Diagnósticos': conjunto_doencas,
        'Localizações': {cuidador.localizacao for cuidador in cuidadores},
        'Escolaridades': {cuidador.escolaridade for cuidador in cuidadores},
        'Referenciações': list(dict.fromkeys({cuidador.referenciacao for cuidador in cuidadores})),
    }

    # Obter campos para filtrar por (COG)
    participantes = Participante.objects.all()
    filtrados_cog = participantes.filter(grupo=None)

    conjunto_doencas = set()
    for participante in participantes:
        conjunto_doencas.update(participante.diagnosticos.all())

    # conjunto_referencias = set()
    # for participante in participantes:
    #     conjunto_referencias.update(set(cuidador.obter_reference))

    lista_pesquisa_participantes = {
        # 'Diagnósticos': list(dict.fromkeys({diagnostico for diagnostico in participante.diagnosticos.all() for participante in participantes})),
        'Diagnósticos': conjunto_doencas,
        'Localizações': list(dict.fromkeys(
            {participante.localizacao for participante in participantes if len(participante.localizacao) > 1})),
        'Escolaridades': list(dict.fromkeys({participante.escolaridade for participante in participantes})),
        'Referenciações': list(dict.fromkeys({participante.referenciacao for participante in participantes})),
        'GDS': list(dict.fromkeys({participante.nivel_gds for participante in participantes})),
    }

    selecoes = {}

    if request.POST:
        #print(request.POST)
        if len(request.POST.get('nome')) > 0:
            g = Grupo(
                nome=request.POST.get('nome'),
                programa=request.POST.get('programa'),
            )
            g.save()
            if request.POST.get('programa') == 'CARE':
                for id in request.POST.get('participantes').split(','):
                    c = Cuidador.objects.get(id=id)
                    g.cuidadores.add(c)

            elif request.POST.get('programa') == 'COG':
                for id in request.POST.get('participantes').split(','):
                    p = Participante.objects.get(id=id)
                    g.participantes.add(p)

            if g.referenciacao_most_frequent is not None:
                g.referenciacao = Reference.objects.get(nome=g.referenciacao_most_frequent)
            if g.diagnostico_most_frequent is not None:
                g.diagnostico = Doenca.objects.get(nome=g.diagnostico_most_frequent)
            g.localizacao= g.localizacao_most_frequent
            g.escolaridade= g.escolaridade_most_frequent
            g.save()


            # Criar as partes e sessoes para este grupo
            for sessao in Sessao.objects.filter(programa=g.programa).all():
                sessao_grupo = SessaoDoGrupo(grupo=g, sessao=sessao)
                sessao_grupo.save()
                if g.programa == 'CARE':
                    for parte in sessao.partes.all():
                        parte_grupo = ParteGrupo.objects.create(
                            sessaoGrupo=sessao_grupo,
                            parte=parte
                        )
                        parte_grupo.save()
                elif g.programa == 'COG':
                    for exercicio in sessao.exercicios.all():
                        parte_grupo = ParteGrupo.objects.create(
                            sessaoGrupo=sessao_grupo,
                            exercicio=exercicio
                        )
                        parte_grupo.save()



    contexto = {
        'tem_proxima': tem_proxima,
        'grupos': Grupo.objects.all(),
        'cuidadores': Cuidador.objects.filter(grupo=None),
        'formGrupo': formGrupo,
        'lista_pesquisa_cuidadores': lista_pesquisa_cuidadores,
        'lista_pesquisa_participantes': lista_pesquisa_participantes,
        'filtrados_care': filtrados_care,
        'filtrados_cog': filtrados_cog,
        'selecoes': selecoes
    }
    return render(request, 'diario/new_group_remake.html', contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def obter_cadidatos(request):
    participantes = None
    if request.method == 'POST':
        #print(request.POST)
        match (request.POST.get('programa')):
            case 'CARE':
                # participantes = Cuidador.objects.filter(grupo=None)
                participantes = Cuidador.objects.all()
                if len(request.POST.get('localizacao')) > 0:
                    participantes = participantes.filter(localizacao=request.POST.get('localizacao'))
                if len(request.POST.get('diagnostico')) > 0:
                    participantes = participantes.filter(participantes__diagnosticos__in=request.POST.get('diagnostico'))
                if len(request.POST.get('escolaridade')) > 0:
                    participantes = participantes.filter(escolaridade=request.POST.get('escolaridade'))
                if len(request.POST.get('referenciacao')) > 0:
                    participantes = participantes.filter(
                        referenciacao=Reference.objects.get(id=request.POST.get('referenciacao')))

            case 'COG':
                # participantes = Participante.objects.filter(grupo=None)
                participantes = Participante.objects.all()
                if len(request.POST.get('localizacao')) > 0:
                    participantes = participantes.filter(localizacao=request.POST.get('localizacao'))
                if len(request.POST.get('diagnostico')) > 0:
                    participantes = participantes.filter(diagnosticos__in=request.POST.get('diagnostico'))
                if len(request.POST.get('escolaridade')) > 0:
                    participantes = participantes.filter(escolaridade=request.POST.get('escolaridade'))
                if len(request.POST.get('referenciacao')) > 0:
                    participantes = participantes.filter(
                        referenciacao=Reference.objects.get(id=request.POST.get('referenciacao')))
                if len(request.POST.get('gds')) > 0:
                    participantes = participantes.filter(nivel_gds=int(request.POST.get('gds')))

    # print(participantes)
    contexto = {
        'programa': request.POST.get('programa'),
        'participantes': participantes,
    }
    return render(request, "diario/obter_candidatos.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def guarda_grupo(request):
    if request.method == 'POST':
        nome = request.POST['nome']

        novo_grupo = Grupo(nome=nome)
        if 'diagnostico' in request.POST:
            novo_grupo.diagnostico = Doenca.objects.get(doenca=request.POST['diagnostico'])
        if 'localizacao' in request.POST:
            novo_grupo.localizacao = request.POST['localizacao']
        if 'escolaridade' in request.POST:
            novo_grupo.escolaridade = request.POST['escolaridade']
        if 'referenciacao' in request.POST:
            novo_grupo.referenciacao = Reference.objects.get(reference=request.POST['referenciacao'])
        if 'programa' in request.POST:
            novo_grupo.programa = request.POST['programa']

        novo_grupo.save()

        # Cria todos os objetos sessaoGrupo e parteGrupo (que registam detalhes das sessoes e partes do grupo)
        for sessao in Sessao.objects.filter(programa=novo_grupo.programa).all():
            sessao_grupo = SessaoDoGrupo(grupo=novo_grupo, sessao=sessao)
            sessao_grupo.save()
            if novo_grupo.programa == 'CARE':
                for parte in sessao.partes.all():
                    parte_grupo = ParteGrupo.objects.create(
                        sessaoGrupo=sessao_grupo,
                        parte=parte
                    )
                    parte_grupo.save()
            elif novo_grupo.programa == 'COG':
                for exercicio in sessao.exercicios.all():
                    parte_grupo = ParteGrupo.objects.create(
                        sessaoGrupo=sessao_grupo,
                        exercicio=exercicio
                    )
                    parte_grupo.save()

        for cuidador_id in request.POST.getlist('cuidadores_selecionados[]'):
            # print(f"cuidador selecionado: {cuidador_id}")
            c = Cuidador.objects.get(id=int(cuidador_id))
            novo_grupo.cuidadores.add(c)

    return redirect('diario:dashboard_Care')


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_group_details(request, grupo_id):
    cuidadores = Cuidador.objects.filter(grupo=grupo_id)
    mentores = Mentor.objects.filter(grupo=grupo_id)
    dinamizadores = DinamizadorConvidado.objects.filter(grupo=grupo_id)

    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)

    print(request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']))
    contexto = {
        'grupo': Grupo.objects.get(id=grupo_id),
        'cuidadores': cuidadores,
        'mentores': mentores,
        'dinamizadores': dinamizadores,
        'grupos_permissoes' : request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
        'tem_proxima': tem_proxima,
    }
    return render(request, "diario/detalhes_grupo.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def group_members(request, grupo_id):
    cuidadores = Cuidador.objects.filter(grupo=grupo_id)
    mentores = Mentor.objects.filter(grupo=grupo_id)
    dinamizadores = DinamizadorConvidado.objects.filter(grupo=grupo_id)

    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)

    # formDinamizador = DinamizadorForm(request.POST or None)
    # if formDinamizador.is_valid():
    #     formDinamizador.save()
    #     return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))

    contexto = {
        'tem_proxima': tem_proxima,
        'grupo_id': grupo_id,
        'grupo': Grupo.objects.get(id=grupo_id),
        'cuidadores': cuidadores,
        'mentores': mentores,
        'dinamizadores': dinamizadores,
        # 'formDinamizador': formDinamizador,
        'dinami': DinamizadorConvidado.objects.filter(grupo=None),
        'caregiver': Cuidador.objects.filter(grupo=None),
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }
    return render(request, "diario/group_members.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def group_sessions(request, grupo_id):
    # agora podemos usar sessao__programa="CARE" ou ="COG" para diferenciar entre os dois programas

    sessoes_do_grupo = SessaoDoGrupo.objects.filter(grupo=grupo_id)
    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)


    sessoes = Sessao.objects.all()
    grupo = Grupo.objects.get(id=grupo_id)
    sessao_em_curso = None
    proxima_sessao = None

    for sessao in sessoes_do_grupo:
        if sessao.estado == 'EC':
            sessao_em_curso = sessao.id
            break
        if sessao.estado == 'PR':
            proxima_sessao = sessao.id
            break

    #    sessoes = Grupo.objects.get(id=grupo_id).sessoes.all()

    contexto = {
        'tem_proxima': tem_proxima,
        'sessoes_do_grupo': sessoes_do_grupo,
        'grupo': grupo,
        'proxima_sessao': proxima_sessao,
        'sessao_em_curso': sessao_em_curso,
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }
    return render(request, "diario/group_sessions.html", contexto)


def group_sessions_cog(request, grupo_id):
    # agora podemos usar sessao__programa="CARE" ou ="COG" para diferenciar entre os dois programas
    sessoes_do_grupo = SessaoDoGrupo.objects.filter(grupo=grupo_id)
    grupo = Grupo.objects.get(id=grupo_id)

    for sessao in sessoes_do_grupo:
        if sessao.estado == 'PR':
            proxima_sessao = sessao.id
            break
    else:
        proxima_sessao = -1

    #    sessoes = Grupo.objects.get(id=grupo_id).sessoes.all()

    contexto = {
        'sessoes_do_grupo': sessoes_do_grupo,
        'grupo': grupo,
        'proxima_sessao': proxima_sessao
    }
    return render(request, "diario/group_sessions.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def group_notes(request, grupo_id):
    contexto = {
        'grupo': Grupo.objects.get(id=grupo_id),

        'notasGrupo': NotaGrupo.objects.filter(grupo=grupo_id),
    }
    return render(request, "diario/group_notes.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def caregiver_update(request, cuidador_id, grupo_id):
    cuidador = Cuidador.objects.get(pk=cuidador_id)
    formCuidador = CuidadorForm(request.POST or None, instance=cuidador)

    if formCuidador.is_valid():
        formCuidador.save()
        return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))

    contexto = {
        'grupo_id': grupo_id,
        'formCuidador': formCuidador,
    }

    return render(request, "diario/caregiver_update.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def create_caregiver(request, grupo_id):
    grupo = Grupo.objects.get(id=grupo_id)

    if request.method == 'POST':
        formCuidador = CuidadorForm(request.POST, request.FILES)
        if formCuidador.is_valid():
            informacao_sensivel = InformacaoSensivel()
            informacao_sensivel.nome = formCuidador.cleaned_data['nome']
            informacao_sensivel.email = formCuidador.cleaned_data['email']
            informacao_sensivel.telemovel = formCuidador.cleaned_data['telemovel']
            informacao_sensivel.save()

            user = User()
            user.username = formCuidador.cleaned_data['username']
            user.password = formCuidador.cleaned_data['password']
            user.email = formCuidador.cleaned_data['email']
            user.save()

            my_group = DjangoGroup.objects.get(name='Cuidador') 
            my_group.user_set.add(new_user)

            cuidador = Cuidador()
            cuidador.user = user
            cuidador.escolaridade = formCuidador.cleaned_data['escolaridade']
            cuidador.nascimento = formCuidador.cleaned_data['nascimento']
            cuidador.nacionalidade = formCuidador.cleaned_data['nacionalidade']
            cuidador.localizacao = formCuidador.cleaned_data['localizacao']
            cuidador.referenciacao = Reference.objects.filter(nome=formCuidador.cleaned_data['referenciacao']).first()
            cuidador.info_sensivel = informacao_sensivel
            cuidador.save()

            cuidador.grupo.add(grupo)
            cuidador.save()

        return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))
    else:
        formCuidador = CuidadorForm()

    contexto = {
        'formCuidador': formCuidador,
        'grupo_id': grupo_id,
    }

    return render(request, "diario/create_caregiver.html", contexto)

@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def create_dinamizador(request, grupo_id):
    grupo = Grupo.objects.get(id=grupo_id)

    if request.method == 'POST':
        formDinamizador = DinamizadorForm(request.POST, request.FILES)
        if formDinamizador.is_valid():
            informacao_sensivel = InformacaoSensivel()
            informacao_sensivel.nome = formDinamizador.cleaned_data['nome']
            informacao_sensivel.email = formDinamizador.cleaned_data['email']
            informacao_sensivel.telemovel = formDinamizador.cleaned_data['telemovel']
            informacao_sensivel.save()

            user = User()
            user.username = formDinamizador.cleaned_data['username']
            user.password = formDinamizador.cleaned_data['password']
            user.email = formDinamizador.cleaned_data['email']
            user.save()

            dinamizador = DinamizadorConvidado()
            dinamizador.user = user
            dinamizador.nacionalidade = formDinamizador.cleaned_data['nacionalidade']
            dinamizador.localizacao = formDinamizador.cleaned_data['localizacao']
            dinamizador.nascimento = formDinamizador.cleaned_data['nascimento']
            dinamizador.funcao = formDinamizador.cleaned_data['funcao']
            dinamizador.info_sensivel = informacao_sensivel
            dinamizador.save()

            dinamizador.grupo.add(grupo)
            dinamizador.save()

        return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))
    else:
        formDinamizador = DinamizadorForm()
    contexto = {
        'formDinamizador': formDinamizador,
        'grupo_id': grupo_id,
    }

    return render(request, "diario/new_dinamizador.html", contexto)

@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def profile_care_view(request, cuidador_id, grupo_id):
    formDocument = Documents_Form(request.POST, request.FILES)

    if formDocument.is_valid():
        formDocument.save()
        return HttpResponseRedirect(reverse('diario:p_view', args=(cuidador_id, grupo_id,)))

    contexto = {
        'cuidador': Cuidador.objects.get(pk=cuidador_id),
        'documents': Documents.objects.filter(cuidador=cuidador_id),
        'grupo': Grupo.objects.get(id=grupo_id),
        'notas': Nota.objects.filter(cuidador=cuidador_id),
        'formDocument': formDocument,
        'cuidador_id': cuidador_id,
        'grupo_id': grupo_id

    }
    return render(request, "diario/profile.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def caregiver_delete(request, cuidador_id, grupo_id):
    cuidador = Cuidador.objects.get(pk=cuidador_id)
    grupo = Grupo.objects.get(pk=grupo_id)
    grupo.cuidadores.remove(cuidador)

    return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def dinamizador_delete(request, dinamizador_id, grupo_id):
    dinamizador = DinamizadorConvidado.objects.get(pk=dinamizador_id)
    dinamizador.delete()

    return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def assign_dinamizador(request, grupo_id, dinamizador_id):
    dinamizador = DinamizadorConvidado.objects.get(id=dinamizador_id)
    grupo = Grupo.objects.get(id=grupo_id)
    grupo.dinamizadores.add(dinamizador)

    return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def assign_caregiver(request, grupo_id, cuidador_id):
    cuidador = Cuidador.objects.get(id=cuidador_id)
    grupo = Grupo.objects.get(id=grupo_id)
    grupo.cuidadores.add(cuidador)

    return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def dinamizador_update(request, dinamizador_id, grupo_id):
    dinamizador = DinamizadorConvidado.objects.get(pk=dinamizador_id)
    formDinamizador = DinamizadorForm(request.POST or None, instance=dinamizador)

    if formDinamizador.is_valid():
        formDinamizador.save()
        return HttpResponseRedirect(reverse('diario:group_members', args=(grupo_id,)))

    contexto = {
        'grupo_id': grupo_id,
        'formDinamizador': formDinamizador,

    }

    return render(request, "diario/dina_update.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def delete_groups(request, grupo_id):
    grupo = Grupo.objects.get(pk=grupo_id)
    grupo.delete()

    return HttpResponseRedirect(reverse('diario:dashboard_Care'))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def update_groups(request, grupo_id):
    grupo = Grupo.objects.get(pk=grupo_id)
    formGrupo = GrupoForm(request.POST or None, instance=grupo)

    if formGrupo.is_valid():
        formGrupo.save()
        return HttpResponseRedirect(reverse('diario:grupo_details', args=(grupo_id,)))

    contexto = {
        'grupo_id': grupo_id,
        'formGrupo': formGrupo
    }

    return render(request, "diario/update_groups.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def filter_group(request, cuidador_id):
    cuidador = Cuidador.objects.get(id=cuidador_id)
    filtrados = []

    grupos = Grupo.objects.all()

    lista_pesquisa = {
        'diagnostico': {(grupo.diagnostico.id, grupo.diagnostico.nome) for grupo in grupos if
                        grupo.diagnostico is not None},
        'localizacao': {grupo.localizacao for grupo in grupos if grupo.localizacao != ''},
        'escolaridade': {grupo.escolaridade for grupo in grupos if grupo.escolaridade != ''},
        'referenciacao': {(grupo.referenciacao.id, grupo.referenciacao) for grupo in grupos if
                          grupo.referenciacao is not None}
    }

    selecoes = {}

    if request.POST:
        filtrados = Grupo.objects.all()
        for campo, valor in request.POST.items():
            if valor != '':
                selecoes[campo] = valor

                if campo == 'diagnostico':
                    doenca = Doenca.objects.get(id=valor)
                    filtrados = filtrados.filter(diagnostico=doenca)
                if campo == 'localizacao':
                    filtrados = filtrados.filter(localizacao=valor)
                if campo == 'escolaridade':
                    filtrados = filtrados.filter(escolaridade=valor)
                if campo == 'referenciacao':
                    referencia = Reference.objects.get(id=valor)
                    filtrados = filtrados.filter(referenciacao=referencia)
    contexto = {
        'cuidador': cuidador,
        'lista_pesquisa': lista_pesquisa,
        'grupos': filtrados,
        'selecoes': selecoes
    }

    return render(request, "diario/filter_groups.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def assign_group(request, grupo_id, cuidador_id):
    cuidador = Cuidador.objects.get(id=cuidador_id)
    grupo = Grupo.objects.get(id=grupo_id)
    grupo.cuidadores.add(cuidador)

    return HttpResponseRedirect(reverse('diario:dashboard_Care'))


def login_care_view(request):
    next = request.GET.get('next')
    if request.method == 'POST':
        next = request.POST.get('next')
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            # print(request.POST)
            next_url = request.POST.get('next')
            if next_url:
                return HttpResponseRedirect(next_url)
            else:
                return redirect('diario:dashboard_Care')

    context = {
        'next': next,
    }

    return render(request, 'diario/login.html', context)


def register_user(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = authenticate(username=username, password=password)
            login(request, user)
            return redirect('diario:dashboard_Care')
    else:
        form = UserCreationForm()

    return render(request, 'diario/register_user.html', {
        'form': form,
    })


def logout_care_view(request):
    return render(request, 'mentha/base.html')


def view_iniciar_sessao(request, sessao_grupo_id):
    sessao_grupo = SessaoDoGrupo.objects.get(id=sessao_grupo_id)
    sessao_grupo.estado = SessaoDoGrupo.EMCURSO
    grupo_id = sessao_grupo.grupo.id
    sessao_grupo.inicio = datetime.utcnow()
    sessao_grupo.save()

    # guardar info das presenças: ir buscar info enviada via formulario. e para cada participante guardar na base de dados atualização do utilizador

    if request.method == 'POST':
        for participante_id, tipo_presenca in request.POST.items():
            if participante_id.isdigit():
                presenca = None
                if sessao_grupo.grupo.programa == "CARE":
                    cuidador = Cuidador.objects.get(id=participante_id)
                    presencas = Presenca.objects.filter(cuidador=cuidador, sessaoDoGrupo=sessao_grupo)
                    if not presencas.exists():
                        # Se não houver nenhuma instância, cria uma nova
                        presenca = Presenca.objects.create(
                            cuidador=cuidador,
                            sessaoDoGrupo=sessao_grupo,
                            tipoSessao=Presenca.CARE,
                        )
                    else:
                        # Se houver pelo menos uma instância, atualiza a primeira encontrada
                        presenca = presencas.first()
                elif sessao_grupo.grupo.programa == "COG":
                    participante = Participante.objects.get(id=participante_id)
                    presencas = Presenca.objects.filter(participante=participante, sessaoDoGrupo=sessao_grupo)
                    if not presencas.exists():
                        # Se não houver nenhuma instância, cria uma nova
                        presenca = Presenca.objects.create(
                            participante=participante,
                            sessaoDoGrupo=sessao_grupo,
                            tipoSessao=Presenca.CARE,
                        )
                    else:
                        # Se houver pelo menos uma instância, atualiza a primeira encontrada
                        presenca = presencas.first()

                if tipo_presenca in ["naoVeio", "n"]:
                    presenca.set_faltou()
                elif tipo_presenca in ["online", "o"]:
                    presenca.set_online()
                else:
                    presenca.set_presencial()

                presenca.save()

    return HttpResponseRedirect(reverse('diario:sessao', args=[sessao_grupo_id, grupo_id]))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_sessao(request, sessao_grupo_id, grupo_id):
    apresentacao = ""
    grupo = Grupo.objects.get(id=grupo_id)
    sessao = SessaoDoGrupo.objects.get(id=sessao_grupo_id, grupo=grupo)

    data = sessao.data

    grupos, sg, is_participante = get_grupos(request.user)
    tem_proxima, datas = get_proxima_sessao(grupos)

    pode_iniciar = False
    if data:
        if data.day == datetime.utcnow().day or sessao.inicio is not None:
            pode_iniciar = True

    partes_grupo = ParteGrupo.objects.filter(sessaoGrupo=sessao_grupo_id)
    for parte in partes_grupo:
        if parte.concluido == False:
            if grupo.programa == "CARE":
                proxima_parte = parte.parte.id
            elif grupo.programa == "COG":
                proxima_parte = parte.exercicio.id
            break
    else:
        proxima_parte = 0
    
    tempo_total_partes = 0
    tempo_total_partes_grupo = 0

    if grupo.programa == "CARE":
        participantes = Cuidador.objects.filter(grupo=grupo_id).order_by('info_sensivel__nome')
        for parte in sessao.sessao.partes.all():
            tempo_total_partes += int(parte.duracao)

        for parte_grupo in partes_grupo:
            tempo_total_partes_grupo += int(parte_grupo.duracao_minutos)

    elif grupo.programa == "COG":
        participantes = Participante.objects.filter(grupo=grupo_id).order_by('info_sensivel__nome')
        for exercicio in sessao.sessao.exercicios.all():
            tempo_total_partes += int(exercicio.duracao)

        for parte_grupo in partes_grupo:
            tempo_total_partes_grupo += int(parte_grupo.duracao_minutos)

    contexto = {
        'parte': sessao.sessao.partes,
        'proxima_parte': proxima_parte,
        'tem_proxima': tem_proxima,
        'sessaoGrupo': sessao,
        'partesGrupo': partes_grupo,
        'participantes': participantes,
        'grupo': Grupo.objects.get(id=sessao.grupo.id),
        'pode_iniciar': pode_iniciar,
        'apresentacao' : apresentacao,
        'tempo_total_partes': tempo_total_partes,
        'cuidador': Cuidador.objects.filter(user=request.user).first(),
        'participante': Participante.objects.filter(user=request.user).first(),
        'tempo_total_partes_grupo': tempo_total_partes_grupo,
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }

    return render(request, 'diario/sessao.html', contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_detalhes_sessao(request, id_sessao_grupo):
    sessao = SessaoDoGrupo.objects.get(sessao=id, grupo=id_sessao_grupo)
    partes_grupo = ParteGrupo.objects.filter(sessaoGrupo=sessao)

    contexto = {
        'id': id,
        'sessaoGrupo': sessao,
        'partesGrupo': partes_grupo,
        'participantes': Cuidador.objects.filter(grupo=sessao.grupo.id),
        'grupo': Grupo.objects.get(id=sessao.grupo.id),

    }
    return render(request, "diario/detalhes_sessao.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_diario(request, idGrupo, idSessao):  # NN: Usar sessao_grupo_id em vez de idSessao
    grupo = Grupo.objects.filter(id=idGrupo).get()
    sessao = Sessao.objects.get(id=idSessao)
    sessaoGrupo = Sessao.objects.get(sessao=sessao, grupo=grupo)
    parte = sessao.partes.all()

    contexto = {
        'participantes': Cuidador.objects.filter(grupo=idGrupo),
        'grupo': Grupo.objects.filter(id=idGrupo),
        'sessao': Sessao.objects.filter(id=idSessao),
        'tipo': parte,
        'sessaoGrupo': sessaoGrupo,
    }

    return render(request, "diario/diario.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_diario_participante(request, idSessaoGrupo, idParticipante):
    sessao_grupo = SessaoDoGrupo.objects.get(pk=idSessaoGrupo)
    programa = sessao_grupo.grupo.programa
    lista_ids_escolhas_multiplas = []
    exercicios = []
    form_list = []
    if programa == "CARE":
        participante = Cuidador.objects.get(pk=idParticipante)
        notas = Nota.objects.filter(cuidador=participante, sessao_grupo=sessao_grupo).order_by('-data')
    elif programa == "COG":
        participante = Participante.objects.get(pk=idParticipante)
        notas = Nota.objects.filter(participante=participante, sessao_grupo=sessao_grupo).order_by('-data')
        # respostas = Resposta.objects.filter(participante=participante, sessao_grupo=sessao_grupo)
        exercicios = sessao_grupo.sessao.exercicios.all()
        form_list = []
        for ex in exercicios:
            for parte in ex.partes_do_exercicio.all():
                for pergunta in parte.perguntas.all():
                    pg = ParteGrupo.objects.filter(sessaoGrupo=sessao_grupo, exercicio=ex).get()
                    initial_data = {}
                    r = Resposta.objects.filter(
                        participante__id=idParticipante,
                        sessao_grupo__id=idSessaoGrupo,
                        pergunta=pergunta,
                        parte_exercicio=parte,
                    )

                    if len(r) > 0:
                        r = r.get()
                        initial_data = {
                            'resposta_escrita': r.resposta_escrita,
                            'certo': r.certo,
                        }
                        if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                            lista_ids_escolhas_multiplas.append(r.resposta_escolha.id)

                    if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                        form = RespostaForm_RespostaEscrita_Dinamizador(None, initial=initial_data)
                    elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                        form = RespostaForm_RespostaSubmetida_Dinamizador(None)
                    if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                        form = None

                    tuplo = (pergunta, parte, pg, form)
                    form_list.append(tuplo)

    if request.method == "POST":
        print(request.POST)
        if request.POST.get('partilha'):
            partilha_text = request.POST.get('partilha')
            id_participante = request.POST.get('participante')
            partilha = Partilha(cuidador=Cuidador.objects.get(pk=id_participante), partilha=partilha_text, sessao_grupo=sessao_grupo)
            partilha.save()

        elif request.POST.get('cuidador'):
            cuidador_id = request.POST.get('cuidador')
            nota_text = request.POST.get('nota')
            nota = Nota(cuidador=Cuidador.objects.get(pk=cuidador_id), nota=nota_text, sessao_grupo=sessao_grupo)
            if DinamizadorConvidado.objects.filter(user=request.user).first():
                nota.anotador_dinamizador = DinamizadorConvidado.objects.get(user=request.user)
            nota.sessao_grupo = sessao_grupo
            nota.save()

        elif request.POST.get('participante'):
            participante_id = request.POST.get('participante')
            nota_text = request.POST.get('nota')
            nota = Nota(participante=Participante.objects.get(pk=participante_id), nota=nota_text, sessao_grupo=sessao_grupo)
            if Mentor.objects.filter(user=request.user).first():
                nota.anotador_mentor = Mentor.objects.get(user=request.user)
            nota.sessao_grupo = sessao_grupo
            nota.save()

    context = {
        'participante_id': idParticipante,
        'participante': participante,
        'notas': notas,
        'partilhas': Partilha.objects.filter(sessao_grupo=sessao_grupo).order_by('-data'),
        'informacoes': Informacoes.objects.filter(participante=idParticipante).order_by('-data'),
        'respostas': 'aaa',
        'notaForm': NotaForm(),
        'partilhaForm': PartilhaForm(),
        'participante': participante,
        'sessaoGrupo': sessao_grupo,
        'exercicios': exercicios,
        'form_list': form_list,
        'lista_ids_escolhas_multiplas': lista_ids_escolhas_multiplas,
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }

    return render(request, "diario/diario_participante.html", context)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_atualiza_presencas_diario(request, idSessaoGrupo):
    sessao_grupo = SessaoDoGrupo.objects.get(pk=idSessaoGrupo)
    grupo = sessao_grupo.grupo
    idGrupo = grupo.id

    nomes = request.POST.getlist("nome")
    valores = request.POST.getlist("valor")
    for participante_id, tipo_presenca in zip(nomes, valores):
        presenca = None
        if sessao_grupo.grupo.programa == "CARE":
            presenca = Presenca.objects.get(cuidador=Cuidador.objects.get(id=participante_id),
                                               sessaoDoGrupo=sessao_grupo)
        elif sessao_grupo.grupo.programa == "COG":
            presenca = Presenca.objects.get(participante=Participante.objects.get(id=participante_id),
                                               sessaoDoGrupo=sessao_grupo)
        if tipo_presenca in ["naoVeio", "n"]:
            presenca.set_faltou()
        elif tipo_presenca in ["online", "o"]:
            presenca.set_online()
        else:
            presenca.set_presencial()

    context = {
        'participantes': Cuidador.objects.filter(grupo=idGrupo).order_by('user'),
        'grupo_id': idGrupo,
        'notasGrupo': NotaGrupo.objects.filter(grupo=idGrupo),
        'partilhas': PartilhaGrupo.objects.filter(grupo=idGrupo),
        'informacoes': Informacoes.objects.all(),
        'respostas': Resposta.objects.all(),
        'notaGrupoForm': NotaGrupoForm(),
        'partilhaGrupoForm': PartilhaGrupoForm()

    }

    return render(request, "diario/diario_participante.html", context)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_diario_grupo(request, idSessaoGrupo):
    sessao_grupo = SessaoDoGrupo.objects.get(id=idSessaoGrupo)
    idGrupo = sessao_grupo.grupo.id
    programa = sessao_grupo.grupo.programa
    if programa == "CARE":
        participantes = Cuidador.objects.filter(grupo=idGrupo).order_by('info_sensivel__nome')
    elif programa == "COG":
        participantes = Participante.objects.filter(grupo=idGrupo).order_by('info_sensivel__nome')

    form_list = []
    form_nota_grupo = NotaGrupoForm(request.POST or None)
    form_partilhas_grupo = PartilhaGrupoForm(request.POST or None)

    if request.method == "POST":
        print(request.POST)
        form = NotaGrupoForm(request.POST or None)
        if request.POST.get('notaGrupo'):
            nota_grupo_text = request.POST.get('notaGrupo')
            nota_grupo = NotaGrupo(notaGrupo=nota_grupo_text, sessao_grupo=sessao_grupo)
            if DinamizadorConvidado.objects.filter(user=request.user).first():
                nota_grupo.anotador_dinamizador = DinamizadorConvidado.objects.get(user=request.user)
            if Mentor.objects.filter(user=request.user).first():
                nota_grupo.anotador_mentor = Mentor.objects.get(user=request.user)
            nota_grupo.grupo = sessao_grupo.grupo
            nota_grupo.save()
        form1 = PartilhaGrupoForm(request.POST or None)
        if request.POST.get('descricao'):
            partilha_text = request.POST.get('descricao')
            partilha = Partilha(partilha=partilha_text, sessao_grupo=sessao_grupo)
            if DinamizadorConvidado.objects.filter(user=request.user).first():
                partilha.partilha_dinamizador = DinamizadorConvidado.objects.get(user=request.user)
            if Mentor.objects.filter(user=request.user).first():
                partilha.partilha_mentor = Mentor.objects.get(user=request.user)
            partilha.grupo = sessao_grupo.grupo
            partilha.save()

    form = RespostasForm(request.POST or None)
    if form.is_valid():
        form.save()

    online_list = []
    presencial_list = []
    faltou_list = []
    for pessoa in participantes:
        if programa == "CARE":
            presenca = Presenca.objects.filter(cuidador=pessoa, sessaoDoGrupo=sessao_grupo)
        elif programa == "COG":
            presenca = Presenca.objects.filter(participante=pessoa, sessaoDoGrupo=sessao_grupo)

        if len(presenca) > 0:
            presenca = presenca.get()
            if presenca is not None and presenca.mode == Presenca.ONLINE:
                online_list.append(int(pessoa.id))
            elif presenca is not None and presenca.mode == Presenca.PRESENT:
                presencial_list.append(int(pessoa.id))
            else:
                faltou_list.append(int(pessoa.id))
    context = {
        'participantes': participantes,
        'grupo_id': idGrupo,
        'grupo': SessaoDoGrupo.objects.get(id=idSessaoGrupo).grupo,
        'sessaoGrupo': sessao_grupo,
        'notasGrupo': NotaGrupo.objects.filter(grupo=idGrupo, sessao_grupo=sessao_grupo),
        'partilhas': Partilha.objects.filter(sessao_grupo=sessao_grupo).order_by('-data'),
        'informacoes': Informacoes.objects.all(),
        # 'respostas': Respostas.objects.all(),
        'notaGrupoForm': NotaGrupoForm(),
        'partilhaGrupoForm': PartilhaGrupoForm(),
        'online_list': online_list,
        'presencial_list': presencial_list,
        'faltou_list': faltou_list,
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }

    return render(request, "diario/diario_grupo.html", context)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_presencas_sessao(request, proxima_id):
    sessao_grupo = SessaoDoGrupo.objects.get(id=proxima_id)

    contexto = {
        'sessao_grupo': sessao_grupo,
        'participantes': sessao_grupo.grupo.participantes_ou_cuidadores,
    }

    return render(request, "diario/presencas_sessao.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_parteDetalhes(request, parte_do_grupo_id, sessaoGrupo_id, idGrupo):
    sg = SessaoDoGrupo.objects.get(sessao=sessaoGrupo_id, grupo=idGrupo)
    programa = sg.grupo.programa
    if programa == "CARE":
        parte = Parte.objects.get(id=parte_do_grupo_id)
        parte_group = ParteGrupo.objects.get(parte_id=parte, sessaoGrupo=sg)
    elif programa == "COG":
        exercicio = Exercicio.objects.get(id=parte_do_grupo_id)
        parte_group = ParteGrupo.objects.get(exercicio=exercicio, sessaoGrupo=sg)



    q = parte.questionarios.all()
    if len(q) > 0:
        q = parte.questionarios.all()[0]
    else:
        q = None

    contexto = {
        'grupo': idGrupo,
        'id': sessaoGrupo_id,
        'sessaoGrupo': sg,
        'parteGrupo': parte_group,
        'parte': parte,
        'atividades': parte.atividades.all(),
        'numero_sessao': sg.sessao.numeroSessao,
        'q': q,
    }
    return render(request, "diario/parteDetalhes.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_parte(request, parte_do_grupo_id, sessaoGrupo_id, estado, proxima_parte):
    respostas_existentes = {}
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo_id)
    participante = Participante.objects.filter(user=request.user)
    if len(participante) > 0:
        participante = participante.get()
    programa = sg.grupo.programa
    
    contexto = {
        'proxima_parte': proxima_parte,
        'estado': estado,
        'sessaoGrupo': sg,
        'participante': participante,
        'grupos_permissoes': request.user.groups.filter(name__in=['Administrador', 'Dinamizador', 'Mentor']),
    }

    parte_group = None
    if programa == "CARE":
        parte = Parte.objects.get(id=parte_do_grupo_id)
        contexto['parte'] = parte
        contexto['dura'] = parte.duracao
        contexto['atividades'] = parte.atividades.all()
        parte_group = ParteGrupo.objects.get(parte_id=parte, sessaoGrupo=sg)
        q = parte.questionarios.all()
        if len(q) > 0:
            q = parte.questionarios.all()[0]
        else:
            q = None
        contexto['q'] = q


    elif programa == "COG":
        exercicio = Exercicio.objects.get(id=parte_do_grupo_id)
        parte_group = ParteGrupo.objects.get(exercicio=exercicio, sessaoGrupo=sg)

        contexto['exercicio'] = exercicio
        contexto['exercicio_partes'] = exercicio.partes_do_exercicio.all().order_by('ordem')
        contexto['dura'] = exercicio.duracao

        respostas_existentes = {}
        lista_ids_escolhas_multiplas = []

        form_list = []
        for parte in exercicio.partes_do_exercicio.all().order_by('ordem'):
            initial_data = {}
            if parte.perguntas.all():
                for pergunta in parte.perguntas.all():
                    r = Resposta.objects.filter(
                        participante__user=request.user,
                        sessao_grupo__id=sessaoGrupo_id,
                        pergunta=pergunta,
                        # parte_grupo__id = parte_do_grupo_id,
                        parte_exercicio=parte,
                    )

                    if len(r) > 0:
                        r = r.get()
                        initial_data = {
                            'resposta_escrita': r.resposta_escrita
                        }
                        if pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                            lista_ids_escolhas_multiplas.append(r.resposta_escolha.id)

                    if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                        form = RespostaForm_RespostaEscrita(None, initial=initial_data)
                    elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                        form = RespostaForm_RespostaSubmetida(None)
                    elif pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
                        form = None

                    tuplo = (pergunta, parte.ordem, form)
                    form_list.append(tuplo)

            contexto['form_list'] = form_list
            contexto['lista_ids_escolhas_multiplas'] = lista_ids_escolhas_multiplas

    contexto['respostas_existentes'] = respostas_existentes
    contexto['parteGrupo'] = parte_group
    contexto['duracao_segundos'] = parte_group.duracao_minutos * 60
    
    if estado != "ver" and estado != "continuar":
        parte_group.inicio = datetime.utcnow()
        parte_group.save()
    
    return render(request, "diario/parte.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def get_respostas_do_participante(request, idSessaoGrupo, idParteGrupo, idParticipante):
    contexto = {}

    exercicio = Exercicio.objects.get(id=idParteGrupo)
    sg = SessaoDoGrupo.objects.get(id=idSessaoGrupo)
    Participante.objects.get(id=idParticipante)
    parte_group = ParteGrupo.objects.get(exercicio=exercicio, sessaoGrupo=sg)

    form_list = []
    for parte in exercicio.partes_do_exercicio.all():
        if parte.perguntas.all():
            for pergunta in parte.perguntas.all():
                if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                    form = RespostaForm_RespostaEscrita(None, initial={'pergunta': pergunta})
                elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                    form = RespostaForm_RespostaSubmetida(None, initial={'pergunta': pergunta})
                tuplo = (pergunta, form)
                form_list.append(tuplo)
        contexto['form_list'] = form_list

    return render(request, "diario/respostas.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_questionario(request, idPergunta, idParte, sessaoGrupo):
    parte = Parte.objects.get(id=idParte)
    questionario = parte.questionarios.all().filter(id=idPergunta).get()

    if questionario.topico == 'Avaliação de satisfação':
        return redirect('diario:questionario_satisfacao', idPergunta=idPergunta, idParte=idParte,
                        sessaoGrupo=sessaoGrupo)

    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo)
    sg_anterior = SessaoDoGrupo.objects.filter(sessao=questionario.continuacaoDe, grupo=sg.grupo)
    if len(sg_anterior) > 0:
        sg_anterior = sg_anterior[0]
        # print(sg_anterior)

    numero_sessao = sg.sessao.numeroSessao
    pg = ParteGrupo.objects.filter(sessaoGrupo=sg, parte=parte).get()
    lista_opcoes = [x.resposta for x in questionario.perguntas.all()[0].opcoes.all()]

    if request.method == 'POST':
        for key in request.POST:
            if 'choice' in str(key):
                k, pergunta_id = key.split('-')
                q = Pergunta.objects.get(pk=pergunta_id)
                opcao = Opcao.objects.get(pk=int(request.POST.get(key)))
                sg = SessaoDoGrupo.objects.get(pk=request.POST.get('sessaogrupo-id'))
                existing = Escolha.objects.filter(pergunta=q, utilizador=request.user, parte_grupo=pg)
                if len(existing) < 1:
                    nova_escolha = Escolha(opcao=opcao, pergunta=q, utilizador=request.user, parte_grupo=pg,
                                           sessao_grupo=sg)
                    nova_escolha.save()
                else:
                    existing = existing[0]
                    existing.opcao = opcao
                    existing.save()
        return redirect('diario:sessao', sg.id, sg.grupo.id)

    contexto = {
        'idPergunta': idPergunta,
        'idParte': idParte,
        'sessaoGrupo': sessaoGrupo,
        'questionario': questionario,
        'parte': parte,
        'escolhas': Escolha.objects.all(),
        'lista_opcoes': lista_opcoes,
        'parteGrupo': pg.id,
        'numero_sessao': numero_sessao,
        'sg_anterior': sg_anterior,
    }
    return render(request, "diario/questionario.html", contexto)


def partilha_parte(request, sessaoGrupo, idParteExercico):
    # print('partilha parte')
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo)
    sg.parte_ativa = Parte_Exercicio.objects.get(id=idParteExercico)
    sg.save()
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)('test', {
        'type': 'chat_message',
        'message': f'{sg.id}',
    })
    return HttpResponse("OK")


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_avaliacao_participantes(request, sessaoGrupoid):
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupoid)

    participantes = sg.grupo.participantes.all()

    obs_part = ""
    obs_sessao = ""

    participantes_form_list = []
    for participante in participantes:
        existente = AvaliacaoParticipante.objects.filter(sessao_grupo=sg, participante=participante,
                                                         submetido_por=Facilitador.objects.get(user=request.user))
        if len(existente) > 0:
            existente = existente.get()
            obs_part = existente.observacao
            initial_data = {
                'interesse': existente.interesse,
                'comunicacao': existente.comunicacao,
                'iniciativa': existente.iniciativa,
                'satisfacao': existente.satisfacao,
                'humor': existente.humor,
                'eficacia_relacional': existente.eficacia_relacional,
            }
            form = AvaliacaoParticipanteForm(None, initial=initial_data)
        else:
            form = AvaliacaoParticipanteForm(None)
        tuplo = (participante, form)
        participantes_form_list.append(tuplo)

    existente = AvaliacaoSessao.objects.filter(sessao_grupo=sg)

    if len(existente) > 0:
        existente = existente.get()
        obs_sessao = existente.observacao
        initial_data = {
            'planificacao_conteudos': existente.planificacao_conteudos,
            'adq_conteudos': existente.adq_conteudos,
            'adq_materiais': existente.adq_materiais,
            'adq_tempo': existente.adq_tempo,
            'grau_dominio': existente.grau_dominio,
            'necessidade_treino': existente.necessidade_treino,
            'apreciacao_global': existente.apreciacao_global,
            'tipo_treino_competencias': existente.tipo_treino_competencias,
        }
        form_sessao = AvaliacaoSessaoForm(None, initial=initial_data)
    else:
        form_sessao = AvaliacaoSessaoForm(None)
    # form_sessao = AvaliacaoSessaoForm(None)

    contexto = {
        'participantes': participantes,
        'request': request,
        # 'form_paciente' : form_paciente,
        'form_sessao': form_sessao,
        'sg_id': sg.id,
        'participantes_form_list': participantes_form_list,
        'obs_part': obs_part,
        'obs_sessao': obs_sessao,
    }

    return render(request, "diario/avaliacao_participante.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def guarda_avaliacao_participante(request, sessaoGrupo_id):
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo_id)
    participante_id = request.POST.get('participante')

    avaliacao_existente = AvaliacaoParticipante.objects.filter(
        sessao_grupo__id=sessaoGrupo_id,
        participante__id=participante_id,
        submetido_por=Facilitador.objects.get(user=request.user)
    )

    # print(request.POST)

    if len(avaliacao_existente) > 0:
        avaliacao_existente = avaliacao_existente.get()
        avaliacao_existente.interesse = request.POST.get('interesse')
        avaliacao_existente.comunicacao = request.POST.get('comunicacao')
        avaliacao_existente.iniciativa = request.POST.get('iniciativa')
        avaliacao_existente.satisfacao = request.POST.get('satisfacao')
        avaliacao_existente.humor = request.POST.get('humor')
        avaliacao_existente.eficacia_relacional = request.POST.get('eficacia_relacional')
        avaliacao_existente.observacao = request.POST.get('observacao')
        avaliacao_existente.save()
    else:
        avaliacao = AvaliacaoParticipante(
            participante=Participante.objects.get(id=request.POST.get('participante')),
            sessao_grupo=sg,
            interesse=request.POST.get('interesse'),
            comunicacao=request.POST.get('comunicacao'),
            iniciativa=request.POST.get('iniciativa'),
            satisfacao=request.POST.get('satisfacao'),
            humor=request.POST.get('humor'),
            eficacia_relacional=request.POST.get('eficacia_relacional'),
            observacao=request.POST.get('observacao'),
            submetido_por=Facilitador.objects.get(user=request.user)
        )
        avaliacao.save()

    return HttpResponse("OK")


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def guarda_avaliacao_sessao(request, sessaoGrupo_id):
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo_id)

    avaliacao_existente = AvaliacaoSessao.objects.filter(
        sessao_grupo__id=sessaoGrupo_id,
        submetido_por=Facilitador.objects.get(user=request.user)
    )

    if len(avaliacao_existente) > 0:
        avaliacao_existente = avaliacao_existente.get()
        avaliacao_existente.planificacao_conteudos = request.POST.get('planificacao_conteudos')
        avaliacao_existente.adq_conteudos = request.POST.get('adq_conteudos')
        avaliacao_existente.adq_materiais = request.POST.get('adq_materiais')
        avaliacao_existente.adq_tempo = request.POST.get('adq_tempo')
        avaliacao_existente.grau_dominio = request.POST.get('grau_dominio')
        avaliacao_existente.necessidade_treino = request.POST.get('necessidade_treino')
        avaliacao_existente.apreciacao_global = request.POST.get('apreciacao_global')
        avaliacao_existente.tipo_treino_competencias = request.POST.get('tipo_treino_competencias')
        avaliacao_existente.observacao = request.POST.get('observacao')
        avaliacao_existente.save()
    else:
        avaliacao = AvaliacaoSessao(
            sessao_grupo=sg,
            planificacao_conteudos=request.POST.get('planificacao_conteudos'),
            adq_conteudos=request.POST.get('adq_conteudos'),
            adq_materiais=request.POST.get('adq_materiais'),
            adq_tempo=request.POST.get('adq_tempo'),
            grau_dominio=request.POST.get('grau_dominio'),
            necessidade_treino=request.POST.get('necessidade_treino'),
            apreciacao_global=request.POST.get('apreciacao_global'),
            tipo_treino_competencias=request.POST.get('tipo_treino_competencias'),
            observacao=request.POST.get('observacao'),
            submetido_por=Facilitador.objects.get(user=request.user),
        )
        avaliacao.save()

    return HttpResponse("OK")


@check_user_able_to_see_page('Todos')
def view_exercicio(request, idExercicio, parteGrupo, sessaoGrupo):
    parte_grupo = ParteGrupo.objects.get(id=parteGrupo)
    parte = parte_grupo.parte
    sessao_grupo = SessaoDoGrupo.objects.get(id=sessaoGrupo)
    exercicio = Exercicio.objects.get(id=idExercicio)

    # if request.method == 'POST':
    #     print('post')

    contexto = {
        'request': request,
        'exercicio': exercicio,
        'parte_grupo': parte_grupo,
        'sessao_grupo': sessao_grupo,

    }
    return render(request, "diario/exercicio.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_questionario_satisfacao(request, idPergunta, idParte, sessaoGrupo):
    parte = Parte.objects.get(id=idParte)
    questionarios_satisfacao = parte.questionarios.all().filter(topico="Avaliação de satisfação")
    questionario_experiencia = Questionario.objects.all().filter(nome="Avaliação da Experiência de Participação").get()
    sg = SessaoDoGrupo.objects.get(id=sessaoGrupo)
    numero_sessao = sg.sessao.numeroSessao
    pg = ParteGrupo.objects.filter(sessaoGrupo=sg, parte=parte).get()

    q_logistica = questionarios_satisfacao.filter(nome="Logística e Organização").get()
    q_expectativas = questionarios_satisfacao.filter(nome="Expectativas").get()
    q_documentacao = questionarios_satisfacao.filter(nome="Documentação").get()
    q_avaliacao_dinamizadores = questionarios_satisfacao.filter(nome="Avaliação dos Dinamizadores").get()
    q_avaliacao_geral = Questionario.objects.all().filter(nome="Avaliacao Geral do Programa").get()

    lista_opcoes_satisfacao = [x.resposta for x in questionarios_satisfacao[0].perguntas.all()[0].opcoes.all()]
    lista_opcoes_experiencia = [x.resposta for x in questionario_experiencia.perguntas.all()[0].opcoes.all()]
    lista_opcoes_geral = [x.resposta for x in q_avaliacao_geral.perguntas.all()[0].opcoes.all()]

    q_avaliacao_geral_text = q_avaliacao_geral.perguntas.all()[0].texto
    q_avaliacao_geral_long_text = q_avaliacao_geral.perguntas.all()[1]

    rl = Escolha.objects.all().filter(utilizador=request.user, pergunta=q_avaliacao_geral_long_text, parte_grupo=pg,
                                      sessao_grupo=sg)
    if len(rl) > 0:
        resposta_longa = rl.get().resposta_escrita
    else:
        resposta_longa = ""

    if request.method == 'POST':
        for key in request.POST:
            if 'choice' in str(key):
                k, pergunta_id = key.split('-')
                q = Pergunta.objects.get(pk=pergunta_id)
                opcao = Opcao.objects.get(pk=int(request.POST.get(key)))
                sg = SessaoDoGrupo.objects.get(pk=request.POST.get('sessaogrupo-id'))
                existing = Escolha.objects.filter(pergunta=q, utilizador=request.user, parte_grupo=pg)
                if len(existing) < 1:
                    nova_escolha = Escolha(opcao=opcao, pergunta=q, utilizador=request.user, parte_grupo=pg,
                                           sessao_grupo=sg)
                    nova_escolha.save()
                else:
                    existing = existing[0]
                    existing.opcao = opcao
                    existing.save()

            elif 'text' in str(key):
                k, pergunta_id = key.split('-')
                q = Pergunta.objects.get(pk=pergunta_id)
                existing = Escolha.objects.filter(pergunta=q, utilizador=request.user, parte_grupo=pg)
                r = request.POST.get(key)
                if len(existing) < 1:
                    nova_escolha = Escolha(pergunta=q, utilizador=request.user, parte_grupo=pg, sessao_grupo=sg,
                                           resposta_escrita=r)
                    nova_escolha.save()
                else:
                    existing = existing[0]
                    existing.resposta_escrita = r
                    existing.save()

    contexto = {
        'idPergunta': idPergunta,
        'idParte': idParte,
        'sessaoGrupo': sessaoGrupo,
        'questionarios_satisfacao': questionarios_satisfacao,
        'questionario': questionarios_satisfacao[0],
        'parte': parte,
        'escolhas': Escolha.objects.all(),
        'lista_opcoes_satisfacao': lista_opcoes_satisfacao,
        'parteGrupo': pg.id,
        'numero_sessao': numero_sessao,
        'q_logistica': q_logistica,
        'q_expectativas': q_expectativas,
        'q_documentacao': q_documentacao,
        'q_avaliacao_dinamizadores': q_avaliacao_dinamizadores,
        'lista_opcoes_experiencia': lista_opcoes_experiencia,
        'questionario_experiencia': questionario_experiencia,
        'q_avaliacao_geral': q_avaliacao_geral,
        'lista_opcoes_geral': lista_opcoes_geral,
        'q_avaliacao_geral_text': q_avaliacao_geral_text,
        'q_avaliacao_geral_long_text': q_avaliacao_geral_long_text,
        'resposta_longa': resposta_longa,
        # 'sg_anterior' : sg_anterior,
    }

    return render(request, "diario/questionario_satisfacao.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_abrirQuestionario(request, idPergunta, idParte, sessaoGrupo):
    factory = qrcode.image.svg.SvgImage
    uri = request.build_absolute_uri('view_questionario')
    uri = uri.replace('abrirQ', 'q')
    uri = uri.replace('view_questionario', f'{sessaoGrupo}')
    img = qrcode.make(uri, image_factory=factory, box_size=20)
    stream = BytesIO()
    img.save(stream)
    parte = Parte.objects.get(id=idParte)
    questionario = parte.questionarios.all().filter(id=idPergunta)

    contexto = {
        'sessaoGrupo': sessaoGrupo,
        'parte': Parte.objects.filter(id=idParte),
        'atividades': parte.atividades.all(),
        'svg': stream.getvalue().decode(),
        'questionario': questionario,
    }
    return render(request, "diario/abrirQuestionario.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_resultados(request, idPergunta, idParte, sessaoGrupo):
    questionario = Questionario.objects.get(id=idPergunta)
    escolhas = []

    for pergunta in questionario.perguntas.all():
        opcoes_respostas = [opcao.resposta for opcao in Pergunta.objects.get(id=pergunta.id).opcoes.all()]
        opcoes = [opcao for opcao in Pergunta.objects.get(id=pergunta.id).opcoes.all()]
        print(escolhas)
        counter_respostas = {}
        for opcao in opcoes:
            counter_respostas[opcao.id] = 0
        for i, escolha in enumerate(Pergunta.objects.get(id=pergunta.id).escolhas.all()):
            counter_respostas[escolha.opcao.id] += 1




    plt.bar(opcoes_respostas, list(counter_respostas.values()))
    plt.ylabel("respostas")
    plt.autoscale()
    max_ = max(list(counter_respostas.values()))
    steps = list(range(max_ + 1))
    plt.yticks(steps)


    fig = plt.gcf()
    plt.close()

    buf = io.BytesIO()
    fig.savefig(buf, format='png')

    buf.seek(0)
    string = base64.b64encode(buf.read())
    grafico = urllib.parse.quote(string)

    context = {'pergunta': pergunta, 'grafico': grafico, 'sessaoGrupo': sessaoGrupo, 'idParte': idParte,
               }

    return render(request, 'diario/resultados.html', context)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def finalizar_parte(request, idParte, sessao_grupo_id, estado):
    sessao_grupo = SessaoDoGrupo.objects.get(id=sessao_grupo_id)
    grupo_id = sessao_grupo.grupo.id
    programa = sessao_grupo.grupo.programa

    if programa == "CARE":
        parte_group = ParteGrupo.objects.get(parte__id=idParte, sessaoGrupo=sessao_grupo)
    else:
        parte_group = ParteGrupo.objects.get(exercicio__id=idParte, sessaoGrupo=sessao_grupo)

    if estado == "finalizar":
        parte_group.fim = datetime.utcnow()
        parte_group.concluido = True
        parte_group.save()

    # if estado == "continuar":
    # parte_group.fim = datetime.now()
    # parte_group.save()

    return HttpResponseRedirect(reverse('diario:sessao', args=[sessao_grupo_id, grupo_id]))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def voltar_parte(request, idParte, sessao_grupo_id, estado):
    parte_group = ParteGrupo.objects.get(parte_id=idParte, sessaoGrupo_id=sessao_grupo_id)

    if estado == "finalizar":
        parte_group.fim = datetime.utcnow()
        parte_group.concluido = True
        parte_group.save()

    # if estado == "continuar":
    # parte_group.fim = datetime.now()
    # parte_group.save()

    return HttpResponseRedirect(reverse('diario:detalhes_sessao', args=[sessao_grupo_id]))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def finalizar_sessao(request, idGrupo, sessao_grupo_id):
    sessao_group = SessaoDoGrupo.objects.get(id=sessao_grupo_id)
    parte_group_final = sessao_group.parteGrupos.all().last()
    for parte in sessao_group.parteGrupos.all():
        parte.fim = datetime.utcnow()
        parte.concluido = True
        parte.save()

    gera_relatorio_questinarios(sessao_group, request)
    gera_relatorio_diario_bordo(sessao_group, request)

    if request.method == 'POST':
        sessao_group.estado = 'R'
        sessao_group.fim = datetime.utcnow()
        sessao_group.concluido = True
        sessao_group.parte_ativa = None
        sessao_group.save()

    return HttpResponseRedirect(reverse('diario:group_sessions', args=[idGrupo]))


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def view_changeDate(request, sessao_id, group_id):
    sessao = SessaoDoGrupo.objects.get(sessao_id=sessao_id, grupo_id=group_id)
    formDataSessao = SessaoDataForm(request.POST or None, instance=sessao)

    if formDataSessao.is_valid():
        formDataSessao.save()
        return HttpResponseRedirect(reverse('diario:group_sessions', args=[group_id]))

    contexto = {
        'formDataSessao': formDataSessao,
        'grupo_id': group_id
    }

    return render(request, "diario/changeDate.html", contexto)


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def guarda_resposta_view(request, sessaoGrupo_id, parteGrupo_id, utilizador_id, pergunta_id, parte_exercicio_id):
    pergunta = Pergunta_Exercicio.objects.get(id=pergunta_id)
    resposta_existente = Resposta.objects.filter(
        pergunta__id=pergunta_id,
        parte_grupo__id=parteGrupo_id,
        sessao_grupo__id=sessaoGrupo_id,
        participante__id=utilizador_id,
        parte_exercicio__id=parte_exercicio_id
    )
    # print(resposta_existente)
    r = None
    if len(resposta_existente) > 0:
        r = resposta_existente[0]
    else:
        r = Resposta(
            participante=Participante.objects.get(id=utilizador_id),
            pergunta=pergunta,
            sessao_grupo=SessaoDoGrupo.objects.get(id=sessaoGrupo_id),
            parte_grupo=ParteGrupo.objects.get(id=parteGrupo_id),
            parte_exercicio=Parte_Exercicio.objects.get(id=parte_exercicio_id)
        )

    if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
        r.resposta_escrita = request.POST.get('resposta_escrita')
    elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
        r.resposta_submetida = request.FILES.get('file')
    elif pergunta.tipo_resposta == "ESCOLHA_MULTIPLA":
        r.resposta_escolha = Opcao.objects.get(id=request.POST.get('choice'))

    if (request.POST.get('certo') == 'true'):
        r.certo = True

    # print(r)
    r.save()
    # print(request.FILES)
    return HttpResponse("OK")


@login_required(login_url='diario:login')
@check_user_able_to_see_page('Todos')
def respostas_view(request, idSessaoGrupo, idParticipante):
    # USADO PARA ATUALIZAR SO A PARTE DAS RESPOSTAS DO DIARIO
    sessao_grupo = SessaoDoGrupo.objects.get(pk=idSessaoGrupo)
    participante = Participante.objects.get(pk=idParticipante)
    exercicios = sessao_grupo.sessao.exercicios.all()

    form_list = []
    for ex in exercicios:
        for parte in ex.partes_do_exercicio.all():
            for pergunta in parte.perguntas.all():
                r = Resposta.objects.filter(
                    participante__id=idParticipante,
                    sessao_grupo__id=idSessaoGrupo,
                    pergunta=pergunta,
                )

                if len(r) > 0:
                    r = r.get()
                    initial_data = {
                        'resposta_escrita': r.resposta_escrita,
                        'certo': r.certo,
                    }

                if pergunta.tipo_resposta == "RESPOSTA_ESCRITA":
                    form = RespostaForm_RespostaEscrita_Dinamizador(None, initial=initial_data)
                elif pergunta.tipo_resposta == "UPLOAD_FOTOGRAFIA":
                    form = RespostaForm_RespostaSubmetida_Dinamizador(None)
                tuplo = (pergunta, form)
                form_list.append(tuplo)

    context = {
        'participante_id': idParticipante,
        'participante': participante,
        'participante': participante,
        'sessaoGrupo': sessao_grupo,
        'exercicios': exercicios,
        'form_list': form_list
    }
    return render(request, "diario/respostas.html", context)


def gera_relatorio_questinarios(sessaoDoGrupo, request):
    perguntas = []
    document = Document()
    counter = 0

    # Cabeçalho
    paragraph = document.add_paragraph(f'Projeto MentHA')

    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(f'Relatório da {sessaoDoGrupo.sessao.nome}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Relatório
    paragraph = document.add_paragraph("O presente relatório tem como objetivo fornecer os resultados detalhados da "
                                       +f"{sessaoDoGrupo.__str__()}" +
                                       " realizada no dia "+ f"{sessaoDoGrupo.data.day}" + "." +
                                       f"{sessaoDoGrupo.data.month}" + "." +
                                       f"{sessaoDoGrupo.data.year}" + " com a participação de várias pessoas. "
                                       "Durante a sessão abordou-se o seguinte tema: " + f"{sessaoDoGrupo.sessao.nome}")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph = document.add_heading(
        f'Presenças na {sessaoDoGrupo.__str__()}:', 2)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    counter_precensas = {'Presencial': 0, 'Online': 0, 'Faltou': 0}
    for presenca in Presenca.objects.filter(sessaoDoGrupo=sessaoDoGrupo):
        if presenca.present:
            if presenca.mode == Presenca.PRESENT:
                counter_precensas['Presencial'] += 1
            if presenca.mode == Presenca.ONLINE:
                counter_precensas['Online'] += 1
        elif presenca.faltou:
            counter_precensas['Faltou'] += 1

    plt.bar(counter_precensas.keys(), list(counter_precensas.values()))
    plt.ylabel("número de pessoas")
    plt.autoscale()
    max_ = max(list(counter_precensas.values()))
    steps = list(range(max_ + 1))
    plt.yticks(steps)

    fig = plt.gcf()
    plt.close()

    buf = io.BytesIO()
    fig.savefig(buf, format='png')
    buf.seek(0)

    # Add the image to the document
    document.add_picture(buf, width=Inches(4))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption = document.add_paragraph("Presenças por modo de assistência", style='Caption')
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    plt.clf()

    document.add_paragraph("Nesta sessão estiveram presentes em modo online um total de " + str(counter_precensas['Online'])
                           + " participantes e em modo presencial " + str(counter_precensas['Presencial'])
                           + " participantes.")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if counter_precensas['Faltou'] > 1:
        document.add_paragraph("Infelizmente, um total de " + str(counter_precensas['Faltou'])
                               + " participantes não poderam estar presentes")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    elif counter_precensas['Faltou'] == 1:
        document.add_paragraph("Infelizmente, um dos participantes não pode estar presente")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    document.add_page_break()
    paragraph = document.add_heading(
        f'Resultados dos questionários na {sessaoDoGrupo.__str__()}:', 2)
    for escolha in Escolha.objects.filter(sessao_grupo=sessaoDoGrupo):
        perguntas.append(escolha.pergunta)

    for i, pergunta in enumerate(perguntas):
        opcoes_respostas = [opcao.resposta for opcao in Pergunta.objects.get(id=pergunta.id).opcoes.all()]
        opcoes = [opcao for opcao in Pergunta.objects.get(id=pergunta.id).opcoes.all()]
        counter_respostas = {}
        for opcao in opcoes:
            counter_respostas[opcao.id] = 0
        for escolha in Pergunta.objects.get(id=pergunta.id).escolhas.filter(sessao_grupo=sessaoDoGrupo):
            counter_respostas[escolha.opcao.id] += 1

        plt.bar(opcoes_respostas, list(counter_respostas.values()))
        plt.ylabel("respostas")
        plt.autoscale()
        max_ = max(list(counter_respostas.values()))
        steps = list(range(max_ + 1))
        plt.yticks(steps)
        p = document.paragraphs[-1]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        plt.tight_layout()
        fig = plt.gcf()
        plt.close()

        buf = io.BytesIO()
        fig.savefig(buf, format='png')
        buf.seek(0)

        # Add the image to the document
        paragraph = document.add_paragraph("O seguinte gráfico mostra os resultados obtidos da seguinte pergunta: \n"
                                           + pergunta.texto)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_picture(buf, width=Inches(4))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        i += 1
        counter += 1
        caption = document.add_paragraph("Pergunta " + str(i) + "- Resultados", style='Caption')
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if counter == 0:
        paragraph = document.add_paragraph('Não existem resultados disponiveis para esta sessão\n')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Assinatura
    paragraph = document.add_paragraph('O dinamizador,')
    paragraph = document.add_paragraph(f'{request.user.username}')

    # Save the Word document
    nome_ficheiro = 'relatorio'+ sessaoDoGrupo.__str__()
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    document.save(docx_path)

    # Convert the Word document to PDF

    pdf_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.pdf')

    pythoncom.CoInitialize()

    word_app = win32.gencache.EnsureDispatch('Word.Application')
    doc = word_app.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)
    doc.Close()
    word_app.Quit()

    # Create a Django File object from the PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = io.BytesIO(f.read())

    # Assign the PDF file to the file field of sessaoDoGrupo
    sessaoDoGrupo.relatorio.save(f'{nome_ficheiro}.pdf', pdf_data)
    sessaoDoGrupo.save()

    # Delete the temporary files
    os.remove(docx_path)
    os.remove(pdf_path)

def gera_relatorio_diario_bordo(sessaoDoGrupo, request):
    document = Document()
    partilhas = Partilha.objects.all().filter(sessao_grupo=sessaoDoGrupo).order_by('-data')
    notas = Nota.objects.all().filter(sessao_grupo=sessaoDoGrupo)
    notas_grupo = NotaGrupo.objects.all().filter(sessao_grupo=sessaoDoGrupo)
    grupo = sessaoDoGrupo.grupo
    partes = ParteGrupo.objects.filter(sessaoGrupo=sessaoDoGrupo)

    # Cabeçalho
    paragraph = document.add_paragraph(f'Projeto MentHA')

    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(f'Relatório da {sessaoDoGrupo.sessao.nome}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Relatório
    paragraph = document.add_paragraph("O presente relatório tem como objetivo fornecer os dados detalhados acerca do Diário de Bordo da "
                                       +f"{sessaoDoGrupo.__str__()}" +
                                       " realizada no dia "+ f"{sessaoDoGrupo.data.day}" + "." +
                                       f"{sessaoDoGrupo.data.month}" + "." +
                                       f"{sessaoDoGrupo.data.year}" + " com a participação de várias pessoas. "
                                       "Durante a sessão abordou-se o seguinte tema: " + f"{sessaoDoGrupo.sessao.nome}")

    if partes.exists():
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph = document.add_heading(
            f'Partes realizadas na {sessaoDoGrupo.__str__()}:', 2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for parte_grupo in partes:
        if grupo.programa == "CARE":
            paragraph = document.add_paragraph(
                "Fase " + dict(parte_grupo.parte.FASE)[parte_grupo.parte.fase] + " - " + parte_grupo.parte.objetivo)
            paragraph = document.add_paragraph(
                "Duração: " + parte_grupo.duracao_em_horas_minutos + " - " + str(parte_grupo.parte.duracao) + " min")
        if grupo.programa == "COG":
            paragraph = document.add_paragraph(
                "Exercício" + str(parte_grupo.exercicio))
            paragraph = document.add_paragraph(
                "Duração: " + parte_grupo.duracao_em_horas_minutos + " - " + str(parte_grupo.exercicio.duracao) + " min")

    if partilhas.exists():
        paragraph = document.add_heading(
            f'Partilhas realizadas na {sessaoDoGrupo.__str__()}:', 2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for partilha in partilhas:
        if partilha.participante:
            paragraph = document.add_paragraph("Às " + partilha.hora_str() + " o participante " +
                                               partilha.participante.nome + " fez a seguinte partilha: \n" + partilha.partilha + ".")
        elif partilha.cuidador:
            paragraph = document.add_paragraph("Às " + partilha.hora_str() + " o cuidador " +
                                               partilha.cuidador.nome + " fez a seguinte partilha: \n" + partilha.partilha + ".")
        elif partilha.partilha_dinamizador:
            paragraph = document.add_paragraph("Às " + partilha.hora_str() + " o dinamizador " +
                                               partilha.partilha_dinamizador.nome + " fez a seguinte partilha: \n" + partilha.partilha + ".")
        elif partilha.partilha_mentor:
            paragraph = document.add_paragraph("Às " + partilha.hora_str() + " o mentor " +
                                               partilha.partilha_mentor.nome + " fez a seguinte partilha: \n" + partilha.partilha + ".")

    if notas_grupo.exists():
        paragraph = document.add_heading(
            f'Notas realizadas sobre o Grupo {grupo.nome}', 2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for nota_grupo in notas_grupo:
        if nota_grupo.anotador_mentor:
            paragraph = document.add_paragraph("Às " + nota_grupo.hora_str() + " o mentor " +
                                               nota_grupo.anotador_mentor.nome + " fez a seguinte nota sobre o grupo: \n" + nota_grupo.notaGrupo + ".")
        if nota_grupo.anotador_dinamizador:
            paragraph = document.add_paragraph("Às " + nota_grupo.hora_str() + " o dinamizador " +
                                               nota_grupo.anotador_dinamizador.nome + " fez a seguinte nota sobre o grupo: \n" + nota_grupo.notaGrupo + ".")

    if notas.exists() and grupo.programa == "CARE":
        paragraph = document.add_heading(
            f'Notas realizadas acerca dos cuidadores', 2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = document.add_paragraph("Dinamizador " + notas.first().anotador_dinamizador.nome + ":")
    elif notas.exists() and grupo.programa == "COG":
        paragraph = document.add_heading(
            f'Notas realizadas acerca dos participantes', 2)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = document.add_paragraph("Mentor " + notas.first().anotador_dinamizador.nome + ":")


    for nota in notas:
        if nota.anotador_dinamizador:
            paragraph = document.add_paragraph("Às " + nota.hora_str() + " o dinamizador fez a seguinte nota acerca do cuidador "
                                               + nota.cuidador.nome + ": \n" + nota.nota + ".")
        if nota.anotador_mentor:
            paragraph = document.add_paragraph("Às " + nota.hora_str() + " o mentor fez a seguinte nota acerca do participante "
                                               + nota.cuidador.nome + ": \n" + nota.nota + ".")

    # Assinatura
    if DinamizadorConvidado.objects.filter(user=request.user):
        paragraph = document.add_paragraph('O dinamizador,')
        paragraph = document.add_paragraph(f'{request.user.username}')
    elif Mentor.objects.filter(user=request.user):
        paragraph = document.add_paragraph('O mentor,')
        paragraph = document.add_paragraph(f'{request.user.username}')


    # Save the Word document
    nome_ficheiro = 'diarioBordo'+ sessaoDoGrupo.__str__()
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    document.save(docx_path)

    # Convert the Word document to PDF

    pdf_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.pdf')

    pythoncom.CoInitialize()

    word_app = win32.gencache.EnsureDispatch('Word.Application')
    doc = word_app.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)
    doc.Close()
    word_app.Quit()

    # Create a Django File object from the PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = io.BytesIO(f.read())

    # Assign the PDF file to the file field of sessaoDoGrupo
    sessaoDoGrupo.diario_bordo.save(f'{nome_ficheiro}.pdf', pdf_data)
    sessaoDoGrupo.save()

    # Delete the temporary files
    os.remove(docx_path)
    os.remove(pdf_path)

@login_required(login_url='diario:login')
@check_user_able_to_see_page('Cuidador')
def user_dashboard(request):

    cuidador = Cuidador.objects.filter(user=request.user).first()

    grupos = cuidador.grupo.all()
    sessao_grupo = SessaoDoGrupo.objects.filter(grupo=grupos.first(), estado='EC').first()

    factory = qrcode.image.svg.SvgImage
    uri = request.build_absolute_uri('zoom')
    uri = uri.replace('abrirZ', 'z')
    img = qrcode.make(uri, image_factory=factory, box_size=5)
    img_pop = qrcode.make(uri, image_factory=factory, box_size=45)
    stream = BytesIO()
    stream_pop = BytesIO()
    img.save(stream)
    img_pop.save(stream_pop)

    context = {
        'grupos': grupos,
        'proxima': sessao_grupo,
        'ss': bool(sessao_grupo),
        'svg': stream.getvalue().decode(),
        'svg_pop': stream_pop.getvalue().decode(),
    }

    return render(request, "diario/user_dashboard.html", context)