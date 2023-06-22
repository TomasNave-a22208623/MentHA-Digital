import io
from platform import python_compiler
from django.http import HttpResponseRedirect, JsonResponse, FileResponse
from django.shortcuts import render, HttpResponse, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from .models import Protocol, Part, Area, Instrument, Dimension, Section, Question, Resolution, Answer, PossibleAnswer,Risk
from django.urls import reverse
from .functions import *
from .forms import *
from diario.models import *
import json
import os
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter 
from xhtml2pdf import pisa

# Other Imports
import plotly.graph_objects as go
import plotly
import pandas as pd
import time
import hashlib
import random
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt


#word Imports to PDF
from docx import Document
import win32com.client as win32
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import pythoncom
import os


# Create your views here.
@login_required(login_url='login')
def dashboard_view(request):
    doctor = request.user
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions}
    return render(request, 'protocolo/dashboard.html', context)


@login_required(login_url='login')
def dashboard_content_view(request):
    doctor = request.user
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions}
    return render(request, 'protocolo/dashboardcontent.html', context)
    

@login_required(login_url='login')
def teste(resquest):
    return render(resquest, 'protocolo/teste.html')

@login_required(login_url='login')
def protocolos_view(request):
    context = {'protocolos': Protocol.objects.all().order_by('order')}
    return render(request, 'protocolo/protocolos.html', context)

@login_required(login_url='login')
def menu_protocolo_view(request):
    return render(request, 'mentha/app-list.html')

@login_required(login_url='login')
def popup_view(request):
    return render(request, 'protocolo/popupparts.html')

# @login_required(login_url='login')
# def risco_view(request):
#     return render(request, 'protocolo/risco.html')

@login_required(login_url='login')
def registo_view(request):
    doctor = request.user
    # participants = Participante.objects.filter(avaliador=doctor)
    if request.method == 'POST':
        new_registo= InformacaoSensivel()
        new_registo.nome = request.POST.get('nome')
        new_registo.email =  request.POST.get('email')
        new_registo.telemovel = request.POST.get('telemovel')
        new_registo.save()

        referefiacaos = Reference.objects.filter(nome = request.POST.get('referenciacao')).get()
        new_participante = Participante()
        new_participante.info_sensivel = new_registo
        new_participante.referenciacao = referefiacaos
        new_participante.avaliador = doctor
        new_participante.nascimento ==  request.POST.get('nascimento')
        new_participante.save()
    # ew_risk.idade = request.POST.get('idade')
    # form = PatientForm(request.POST or None)
    # avaliador = request.user

    # if request.method == "POST":
    #     form = PatientForm(request.POST)
    #     if form.is_valid():
    #         obj = form.save(commit=False)
    #         obj.save()
    # context ={'form' : form, 'avaliador': avaliador}

    context = {}





    return render(request, 'protocolo/participantes_registo.html')

@login_required(login_url='login')
def incrementar(request, part_id, participant_id):
    part = Part.objects.get(pk = part_id)
    particitant = Participante.objects.get(pk = participant_id)
    nome = part.nome + datetime.now
    partParticipant = ParteDoUtilizador(part=part, participante = particitant, nome = nome)    

    return render(request,'protocolo/parts.html')

@login_required(login_url='login')
def participant_risk_view(request,protocol_id):
    doctor = request.user
    protocolo = Protocol.objects.get(pk=protocol_id)
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions, 'protocolo': protocolo}
    return render(request,'protocolo/protocolo_participants_risk.html',context)
#funcao view para adicionar uma parte a um participante

@login_required(login_url='login')
def nova_pagina_risk_report(request):
    form_report_risk = Risk.objects.all()
    return render(request,'protocolo/visualizacao_risk.html',{'form_report_risk':form_report_risk})

@login_required(login_url='login')
def parte_do_utilizador_add_view(request):
    print('teste1231222')
    if request.method == 'POST':
        partId = request.POST.get('partId')
        print(partId)   
        part= Part.objects.get(pk=partId)
        print(part)
        print('teste12312412312')
        patient =  Participante.objects.get(pk=request.POST.get('patientId'))
        parteDoUtilizador = ParteDoUtilizador(part=part,participante=patient)
        parteDoUtilizador.save()

    
    context = {
        'parteDoUtilizador':parteDoUtilizador,
        'part':part,

        }
    return HttpResponse('200')    
    # return JsonResponse({}, status=201)

@login_required(login_url='login')
def parte_view(request):
    parte = Part.objects.all()#participante = patient)
    #patient = Participante.objects.get(pk=patient_id)
    botao_clicado = int(request.GET.get('botao_clicado ', 0))
    context = {
               'parte':parte,
                'botao_clicado': botao_clicado,
                #'patient':patient
               }
    return render(request, 'protocolo/parts.html', context)

# @login_required(login_url='login')
# def risk_form_view(request):
#     submitted = False
#     form_risk = FormRisk()#request.POST or None)
#     if request.method == "POST":
#         form_risk = FormRisk(request.POST or None)
#         if form_risk.is_valid():
#             form_risk.save()
#             return HttpResponseRedirect('protocolo/parts_risk?submitted=True')
#     else:
#         #form_risk = FormRisk()
#         if 'submitted' in request.GET:
#             submitted = True
#     context = {'form_risk':form_risk, 'submitted':submitted,}
#     return render(request,'protocolo/parts_risk.html',context)

@login_required(login_url='login')
def parts_view(request, protocol_id, patient_id):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    resolutions = Resolution.objects.filter(doctor=request.user, patient=Participante.objects.filter(
        pk=patient_id).get())  # Mudar request.user para o patient depois
    patient = Participante.objects.get(pk=patient_id)
    parts = ParteDoUtilizador.objects.filter(participante = patient)
    parte = Part.objects.all()
    # parte_parte = ParteDoUtilizador.objects.all()
    risk_area = Area.objects.get(id = 47)
    pergunta_risk = Question.objects.get(id = 189)
    print(pergunta_risk)
    # statistics
    answered_list = []
    percentage_list = []
    # data = parte_parte.get().data
    for parteDoUtilizador in parts:
        resolution = resolutions.filter(part=parteDoUtilizador)
        if not resolution:
            answered_list.append(0)
            percentage_list.append(0)
        else:
            s = resolution.get().statistics
            answered_list.append(s.get('total_answered'))
            percentage_list.append(s.get('total_percentage'))
    
    # funcao para apagar partes duplicadas com a mesma data
    # for part in range(len(parts)):
    #     current_part = parts[part]
    #     if part+1 < len(parts): 
    #         next_part = parts[part+1]
    #         if current_part.part.name == next_part.part.name and current_part.data == next_part.data:
    #             next_part.delete()
    #             print('deleted')
    #     else:
    #         break
    # print(answered_list)
    # print(percentage_list)
    end = time.time()
    print("Parts", (end - start))
    context = {'parts': zip(parts, answered_list, percentage_list),
               'parte':parte,
               'risk_area':risk_area,
                'protocol': protocol, 'resolutions': resolutions,
               'patient': patient,
               'pergunta_risk':pergunta_risk,
               }
    return render(request, 'protocolo/parts.html', context)

@login_required(login_url='login')
def areas_view(request, protocol_id, part_id, patient_id):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    part = parteDoUtilizador.part
    areas = Area.objects.filter(part=part).order_by('order')
    patient = Participante.objects.get(pk=patient_id)
    rel_q = Question.objects.filter(name="Relação com o Avaliador").get()
    coop_q = Question.objects.filter(name="Cooperação dada na entrevista").get()
    qs_q = Question.objects.filter(name="Questionário Sociodemográfico").get()
    
    # print(coop_q)

    # ESTOU A CRIAR A RESOLUÇAO AQUI, MAS DEPOIS MUDAR DE SITIO
    r = None
    if Resolution.objects.filter(patient=patient, part=parteDoUtilizador, doctor=request.user).exists():
        r = Resolution.objects.filter(patient=patient, part=parteDoUtilizador, doctor=request.user).get()
    else:
        r = Resolution(patient=patient, part=parteDoUtilizador, doctor=request.user)
        r.initialize_statistics()
        r.save()
        # Sem esta ultima linha a página das àreas vinha vazia
        r = Resolution.objects.get(patient=patient, part=parteDoUtilizador, doctor=request.user)

    # statistics
    # print_nested_dict(r.statistics, 0)
    answered_list = []
    percentage_list = []
    s = r.statistics
    for area in areas:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get('percentage'))
    end = time.time()
    print("Parts", (end - start))
    context = {'areas': zip(areas, answered_list, percentage_list), 'part': parteDoUtilizador, 'protocol': protocol,
               'resolution': r.id, 'patient': patient, 'coop': coop_q, 'rel': rel_q, 'qs': qs_q }
    return render(request, 'protocolo/areas.html', context)


@login_required(login_url='login')
def instruments_view(request, protocol_id, part_id, area_id, patient_id):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    patient = Participante.objects.get(pk=patient_id)

    instruments = Instrument.objects.filter(area=area_id).order_by('order')

    # statistics
    r = Resolution.objects.get(patient=patient, doctor=request.user, part=parteDoUtilizador)
    # print_nested_dict(r.statistics, 0)
    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for instrument in instruments:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get('percentage'))
            quotation_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get('quotation'))
    # print(answered_list)
    # print(percentage_list)
    end = time.time()
    print("Instruments", (end - start))

    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'protocol_id': protocol_id, 'part_id': part_id,
               'area_id': area_id, 'patient_id': patient_id,
               'instruments': [e for e in zip(instruments, answered_list, percentage_list, quotation_list)],
               # 'instruments': zip(instruments, answered_list, percentage_list, quotation_list),
               'resolution': r.id,
               'patient': patient, }

    renderizado = render(request, 'protocolo/instruments.html', context)
    end = time.time()
    print("Instruments", (end - start))
    return renderizado


@login_required(login_url='login')
def dimensions_view(request, protocol_id, part_id, area_id, instrument_id, patient_id):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimensions = Dimension.objects.filter(instrument=instrument_id).order_by('order')
    patient = Participante.objects.get(pk=patient_id)

    if len(dimensions) == 1:
        return redirect('protocolo:sections', protocol_id, part_id, area_id, instrument_id, dimensions.get().id, patient_id)

    # statistics
    r = Resolution.objects.get(patient=patient, doctor=request.user, part=parteDoUtilizador)
    # print_nested_dict(r.statistics.get(str(area.id)).get(str(instrument.id)))
    # print(r.statistics.get(str(area.id)).get(str(instrument.id)).get('9').get('quotation'))

    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for dimension in dimensions:
        if s.get(f'{area.id}') is not None:
            answered_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get('answered'))
            percentage_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get('percentage'))
            quotation_list.append(s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get('quotation'))
    # print(answered_list)
    # print(percentage_list)
    end = time.time()
    print("Dimensions", (end - start))
    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument,
               'dimensions': zip(dimensions, answered_list, percentage_list, quotation_list), 'resolution': r.id,
               'patient': patient, }
    return render(request, 'protocolo/dimensions.html', context)


@login_required(login_url='login')
def sections_view(request, protocol_id, part_id, area_id, instrument_id, dimension_id, patient_id):
    start = time.time()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimension = Dimension.objects.get(pk=dimension_id)
    patient = Participante.objects.get(pk=patient_id)

    sections = Section.objects.filter(dimension=dimension_id).order_by('order')
    if len(sections) == 1:
        return redirect('protocolo:question', protocol_id, part_id, area_id, instrument_id, dimension_id, sections.get().id,
                        patient_id)

    # statistics
    r = Resolution.objects.get(patient=patient, doctor=request.user, part=parteDoUtilizador)
    # print_nested_dict(r.statistics, 0)
    answered_list = []
    percentage_list = []
    quotation_list = []
    s = r.statistics
    for section in sections:
        if s.get(f'{area.id}') is not None:
            answered_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get('answered'))
            percentage_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get(
                    'percentage'))
            quotation_list.append(
                s.get(f'{area.id}').get(f'{instrument.id}').get(f'{dimension.id}').get(f'{section.id}').get(
                    'quotation'))

    # print(answered_list)
    # print(percentage_list)

    if len(sections) == 1:
        return redirect(question_view)
    end = time.time()
    print("Sections", (end - start))
    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument, 'dimension': dimension,
               'sections': zip(sections, answered_list, percentage_list, quotation_list), 'resolution': r.id,
               'patient': patient, }

    # print(instrument.number_of_dimensions)

    return render(request, 'protocolo/sections.html', context)


@login_required(login_url='login')
def question_view(request, protocol_id, part_id, area_id, instrument_id, dimension_id, section_id, patient_id):
    start = time.time()
    doencas = Doenca.objects.all()
    protocol = Protocol.objects.get(pk=protocol_id)
    parteDoUtilizador = ParteDoUtilizador.objects.get(pk=part_id)
    area = Area.objects.get(pk=area_id)
    instrument = Instrument.objects.get(pk=instrument_id)
    dimension = Dimension.objects.get(pk=dimension_id)
    section = Section.objects.get(pk=section_id)
    question = Question.objects.filter(section=section).first()
    form = uploadAnswerForm(request.POST or None)
    form_risk = FormRisk(request.POST or None)
    patient = Participante.objects.get(pk=patient_id)
    r = Resolution.objects.filter(patient=patient, doctor=request.user, part=parteDoUtilizador)
    print('VOU DAR UMA PRINT')
    print(parteDoUtilizador)
    print('VOU DAR UMA PRINT')
    
    if len(r) == 0 and parteDoUtilizador.part.name == 'MentHA-Risk':
        r = Resolution.objects.create(patient=patient, doctor=request.user, part=parteDoUtilizador)
        r.initialize_statistics()
        r.save()
    else:
        r = r.get()
    answers = Answer.objects.filter(resolution=r)
    # print(question.section.number_of_questions)
    # print(question.section.dimension.number_of_questions)
    context = {'area': area, 'part': parteDoUtilizador, 'protocol': protocol, 'instrument': instrument, 'dimension': dimension,
               'section': section, 'question': question, 'form': form, 'resolution': r.id, 'answers': answers,
               'patient': patient, 'doencas': doencas,'form_risk':form_risk
               }
    # print(answers)

    if question.question_type == 10:
        if len(answers.filter(question=question)) > 0:
            a = answers.filter(question=question).get()
            json_answer = json.loads(a.text_answer)
            context['Sequenciacao'] = json_answer.get('Sequenciacao')
            context['Perseverativos'] = json_answer.get('Perseverativos')
            context['Proximidade'] = json_answer.get('Proximidade')
            context['Tempo'] = json_answer.get('Tempo')

    if question.question_type == 3:
        question_list = []
        answered_ids = []
        for question in Question.objects.filter(section=section.id):
            question_list.append(question)
            for answer in answers:
                if question == answer.question:
                    answered_ids.append(question.id)

        # Esta parte permite dividir o question type 3 em dois
        # Um que tem sempre as mesmas respostas, e mostrará a página como uma tabela
        # Outro que as respostas são diferentes e mostrará varias perguntas individualmente, todas na mesma página
        ans = 0
        equal = True
        qset = []
        for question in Question.objects.filter(section=section.id):
            if ans == 0:
                ans = 1
                qset = question.possible_answer_name_list  # print(qset)
            else:
                # print(f"{question.possible_answer_name_list} vs {qset}")
                equal = question.possible_answer_name_list == qset

        context['equal_answers'] = equal
        context['question_list'] = question_list

    existing_answer_id = []
    for answer in answers:
        if answer.resolution == r:
            if answer.question == question:
                if answer.multiple_choice_answer is not None:
                    existing_answer_id.append(answer.multiple_choice_answer.id)
                    context['existing_answer_id'] = existing_answer_id
                if answer.text_answer is not None:
                    form.initial.update({'text_answer': answer.text_answer})
                if answer.submitted_answer:
                    context['submitted_answer'] = answer.submitted_answer.url
                if answer.notes is not None:
                    context['notes'] = answer.notes
                if answer.quotation is not None:
                    context['quotation'] = answer.quotation
                if len(answer.MCCAnswer.all()) >= 1:
                    list = []
                    for mca in answer.MCCAnswer.all():
                        list.append(mca.choice.id)
                    existing_answer_id = list
                    context['existing_answer_id'] = list
    existing_risk = None
    for risk in Risk.objects.all():
        if risk.parteDoUtilizador == parteDoUtilizador:
            existing_risk = risk
    context['existing_risk'] = existing_risk
    if request.method == 'POST':
        existing_answer = None
       
        for answer in answers:
            if answer.question == question:
                existing_answer = answer

        if question.question_type == 1 or question.question_type == 9:
            id_answer = request.POST.get("choice")
            r = Resolution.objects.get(part=parteDoUtilizador, patient=patient, doctor=request.user)
            if existing_answer is None:
                # cria uma nova associação
                new_answer = Answer(question=question, multiple_choice_answer=PossibleAnswer.objects.get(pk=id_answer),
                                    resolution=r)
                quotation = new_answer.multiple_choice_answer.quotation
                new_answer.quotation = quotation
                new_answer.notes = request.POST.get('notes')
                new_answer.save()
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', quotation)
            else:
                # modifica a associação existente
                existing_answer.multiple_choice_answer = PossibleAnswer.objects.get(pk=id_answer)
                quotation = existing_answer.multiple_choice_answer.quotation
                existing_answer.quotation = quotation
                existing_answer.notes = request.POST.get('notes')
                existing_answer.save()
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', quotation)


        elif question.question_type == 2:   
            form = uploadAnswerForm(request.POST, files=request.FILES)
            if form.is_valid():
                new_answer = Answer()
                new_answer.submitted_answer = form.cleaned_data['submitted_answer']
                new_answer.text_answer = form.cleaned_data['text_answer']
                new_answer.quotation = form.cleaned_data['quotation']
                new_answer.notes = form.cleaned_data['notes']
                new_answer.question = question
                new_answer.resolution = r

                if existing_answer is None:
                    # cria uma nova resposta
                    r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                           f'{section_id}')
                    new_answer.resolution = r
                    new_answer.save()
                else:
                    # modifica a resposta existente
                    existing_answer.submitted_answer = new_answer.submitted_answer
                    existing_answer.text_answer = new_answer.text_answer
                    existing_answer.quotation = new_answer.quotation
                    existing_answer.notes = new_answer.notes
                    existing_answer.save()

                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                   new_answer.quotation)

        elif question.question_type == 3:
            for key in request.POST:
                if 'choice' in str(key):
                    k, question_id = key.split('-')
                    q = Question.objects.get(pk=question_id)
                    existing_answers_list = []

                    if q.id in answered_ids:
                        a = Answer.objects.filter(resolution=r, question=q).get()
                        # modifica a associação existente
                        a.multiple_choice_answer = PossibleAnswer.objects.get(pk=request.POST.get(key))
                        quotation = a.multiple_choice_answer.quotation
                        a.quotation = quotation
                        a.notes = request.POST.get('notes')
                        a.save()
                        r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                           quotation)
                    else:
                        a = Answer()
                        a.resolution = r
                        a.multiple_choice_answer = PossibleAnswer.objects.get(pk=request.POST.get(key))
                        a.question = q
                        quotation = a.multiple_choice_answer.quotation
                        a.quotation = quotation
                        a.notes = request.POST.get('notes')
                        a.save()
                        r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                               f'{section_id}')
                        r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                                           quotation)

        elif question.question_type == 4 or question.question_type == 6 or question.question_type == 7 or question.question_type == 8:
            existing = False
            if len(Answer.objects.filter(question=question, resolution=r)) >= 1:
                a = Answer.objects.filter(question=question, resolution=r).get()
                a.delete()
                existing = True
            a = Answer()
            a.resolution = r
            a.question = question
            a.notes = request.POST.get('notes')
            q = 0
            a.save()
            for id in request.POST.getlist("choice"):
                c = MultipleChoicesCheckbox()
                c.answer = a
                pa = PossibleAnswer.objects.filter(pk=id).get()
                c.choice = pa
                q = q + c.choice.quotation
                c.save()

            if 'Repetição I' in question.name:
                l = len(request.POST.getlist("choice"))
                if l == 4:
                    q = 2
                elif l == 3:
                    q = 1
                else:
                    q = 0

            a.quotation = q
            a.save()

            if existing:
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)
            else:
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)

        elif question.question_type == 5:
            # print(request.POST)
            existing = False
            if len(Answer.objects.filter(question=question, resolution=r)) >= 1:
                a = Answer.objects.filter(question=question, resolution=r).get()
                a.delete()
                existing = True
            a = Answer()
            a.resolution = r
            a.question = question
            a.notes = request.POST.get('notes')
            a.save()
            counter = 0
            for i in range(1, 5):
                text_area = request.POST.get(str(i))
                if len(text_area) > 1:
                    counter += len(text_area.split(","))
                    # print(text_area)
                    # print(counter)
                    ti = TextInputAnswer()
                    ti.answer = a
                    ti.seconds = i
                    ti.text = text_area
                    ti.save()
            q = calculate_timer_quotation(question, counter)
            a.quotation = q
            a.save()

            if existing:
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)
            else:
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}', q)

        elif question.question_type == 10:
            erro_sequenciacao = request.POST.get('sequenciacao')
            erro_perseverativos = request.POST.get('perseverativos')
            erro_proximidade = request.POST.get('proximidade')
            tempo = request.POST.get('tempo')
            # print(erro_proximidade, erro_sequenciacao, erro_perseverativos, tempo)
            json_answer = {'Sequenciacao': erro_sequenciacao,
                           'Perseverativos': erro_perseverativos,
                           'Proximidade': erro_proximidade,
                           'Tempo': tempo
                           }

            new_answer = Answer()
            new_answer.question = question
            new_answer.resolution = r
            new_answer.text_answer = json.dumps(json_answer)
            new_answer.quotation = int(tempo) - (
                    int(erro_sequenciacao) + int(erro_perseverativos) + int(erro_proximidade))

            if existing_answer is None:
                # cria uma nova resposta
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                new_answer.save()
            else:
                # modifica a resposta existente
                existing_answer.submitted_answer = new_answer.submitted_answer
                existing_answer.text_answer = new_answer.text_answer
                existing_answer.quotation = new_answer.quotation
                existing_answer.notes = new_answer.notes
                existing_answer.save()

            r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                               new_answer.quotation)

        elif question.question_type == 11:
            print(request.POST)
            if request.POST.get('sexo'):
                patient.sexo = request.POST.get('sexo')
            if request.POST.get('nacionalidade'):
                patient.nacionalidadde = request.POST.get('nacionalidade')
            if request.POST.get('nascimento'):
                patient.nascimento = request.POST.get('nascimento')
            if request.POST.get('escolaridade'):
                patient.escolaridade = request.POST.get('escolaridade')
            if request.POST.get('residencia'):
                patient.residencia = request.POST.get('residencia')
            if request.POST.get('laboral'):
                patient.situacaoLaboral = request.POST.get('laboral')
            if request.POST.get('profissao'):
                patient.profissaoPrincipal = request.POST.get('profissao')
            if request.POST.get('economina'):
                patient.situacaoEconomica = request.POST.get('economica')
            if request.POST.get('civil'):
                patient.estadoCivil = request.POST.get('civil')
            if request.POST.get('agregado'):
                patient.agregadoFamiliar = request.POST.get('agregado')
            if request.POST.get('temFilho'):
                patient.temFilhos = request.POST.get('filhos') == 'sim'
            if request.POST.get('nrFilhos'):
                patient.nrFilhos = int(request.POST.get('nrFilhos'))
            if request.POST.get('saude'):
                patient.autoAvaliacaoEstadoSaude = request.POST.get('saude')
            if request.POST.get('doenca') is not None:
                patient.diagnosticos.clear()
                for id in request.POST.get('doenca'):
                    patient.diagnosticos.add(Doenca.objects.get(id=id))
                patient.save()

        elif question.question_type == 12:
            ris = None
            cwd = os.getcwd()
            cwd2 = os.path.join(cwd, 'protocolo', 'static', 'protocolo', 'data_risk')
            file_path_men = os.path.join(cwd2, 'risk_men.json')
            file_path_women = os.path.join(cwd2, 'risk_women.json')
            new_risk= Risk()
            colestrol_virgula = 0.0
            new_risk.idade = request.POST.get('idade')
            new_risk.sexo = request.POST.get('sexo')
            new_risk.peso = request.POST.get('peso')
            new_risk.altura = request.POST.get('altura')
            new_risk.pressao_arterial = request.POST.get('pressao_arterial')
            new_risk.colestrol_total = request.POST.get('colestrol_total') 
            new_risk.colestrol_hdl = request.POST.get('colestrol_hdl')
            new_risk.colestrol_nao_hdl = request.POST.get('colestrol_nao_hdl')
            new_risk.fumador = request.POST.get('fumador')
            new_risk.diabetes = request.POST.get('diabetes')
            new_risk.hemoglobina_gliciada = request.POST.get('hemoglobina_gliciada')
            new_risk.anos_diabetes = request.POST.get('anos_diabetes')
            new_risk.avc = request.POST.get('avc')
            new_risk.enfarte = request.POST.get('enfarte')
            new_risk.doenca_rins = request.POST.get('doenca_rins')
            new_risk.doenca_pernas = request.POST.get('doenca_pernas')
            new_risk.hipercolestrol = request.POST.get('hipercolestrol')
            new_risk.comentario = request.POST.get('comentario')
            colestrol_virgula = new_risk.colestrol_total
            colestrol_virgula = str(colestrol_virgula)
            if request.POST.get('colestrol_total')[1] == ',':
                print("ENTROU NO IF")
                colestrol_virgula = colestrol_virgula.replace(',','.')
                print(colestrol_virgula)
                print("SAIU DO IF")
            elif request.POST.get('colestrol_total')[1] == '.':
                colestrol_virgula = colestrol_virgula
            if request.POST.get('sexo') == 'F':
                risco = risk_json(file_path_women, new_risk.fumador, new_risk.idade, colestrol_virgula,new_risk.pressao_arterial)
                new_risk.risco_de_enfarte = risco
            elif request.POST.get('sexo') == 'M':
                risco = risk_json(file_path_men, new_risk.fumador, new_risk.idade, colestrol_virgula,new_risk.pressao_arterial)
                new_risk.risco_de_enfarte = risco
            new_risk.parteDoUtilizador = parteDoUtilizador
            new_risk.concluido = True
            r.part.concluido = True
            r.save()
            new_risk.imc = calcular_imc(new_risk.peso, new_risk.altura)
            
            if existing_risk is None:
                # cria uma nova resposta
                r.increment_statistics(f'{part_id}', f'{area_id}', f'{instrument_id}', f'{dimension_id}',
                                       f'{section_id}')
                new_risk.save()
                ris = new_risk
            else:
                # modifica a resposta existente
                existing_risk.idade = request.POST.get('idade')
                existing_risk.sexo = request.POST.get('sexo')
                existing_risk.peso = request.POST.get('peso')
                existing_risk.altura = request.POST.get('altura')
                existing_risk.pressao_arterial = request.POST.get('pressao_arterial')
                existing_risk.colestrol_total = request.POST.get('colestrol_total')
                existing_risk.colestrol_hdl = request.POST.get('colestrol_hdl')
                existing_risk.colestrol_nao_hdl = request.POST.get('colestrol_nao_hdl') 
                existing_risk.fumador = request.POST.get('fumador')
                existing_risk.diabetes = request.POST.get('diabetes')
                existing_risk.hemoglobina_gliciada = request.POST.get('hemoglobina_gliciada')
                existing_risk.anos_diabetes = request.POST.get('anos_diabetes')
                existing_risk.avc = request.POST.get('avc')
                existing_risk.enfarte = request.POST.get('enfarte')
                existing_risk.doenca_rins = request.POST.get('doenca_rins')
                existing_risk.doenca_pernas = request.POST.get('doenca_pernas')
                existing_risk.hipercolestrol = request.POST.get('hipercolestrol')
                existing_risk.comentario = request.POST.get('comentario') 
                existing_risk.imc = calcular_imc(existing_risk.peso, existing_risk.altura)
                if request.POST.get('sexo') == 'F':
                    risco = risk_json(file_path_women, existing_risk.fumador, existing_risk.idade, existing_risk.colestrol_total,existing_risk.pressao_arterial)
                    existing_risk.risco_de_enfarte = risco
                elif request.POST.get('sexo') == 'M':
                    risco = risk_json(file_path_men, existing_risk.fumador, existing_risk.idade, existing_risk.colestrol_total,existing_risk.pressao_arterial)
                    existing_risk.risco_de_enfarte = risco
                existing_risk.save()
                ris = existing_risk
            r.change_quotation(f'{area_id}', f'{instrument_id}', f'{dimension_id}', f'{section_id}',
                               new_risk.risco_de_enfarte)
            username = request.user.username

            gera_relatorio_risk_pdf(ris, patient, username)
            
        if question.question_type == 3:
            return redirect('protocolo:instruments', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            patient_id=patient_id)
        elif question.name == "Relação com o Avaliador" or question.name == "Cooperação dada na entrevista" or question.name == "Questionário Sociodemográfico":
            return redirect('protocolo:areas', protocol_id=protocol_id, part_id=part_id, patient_id=patient_id)
        elif question.section.dimension.name == "None" and question.section.name == "None":
            return redirect('protocolo:instruments', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            patient_id=patient_id)
        elif question.section.name == "None" or question.question_type == 9 or question.section.dimension.number_of_questions == 1:
            return redirect('protocolo:dimensions', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            instrument_id=instrument_id, patient_id=patient_id)
        else:
            return redirect('protocolo:sections', protocol_id=protocol_id, part_id=part_id, area_id=area_id,
                            instrument_id=instrument_id, dimension_id=dimension_id, patient_id=patient_id)
    
    
    end = time.time()
    print("Question", (end - start))
    return render(request, 'protocolo/question.html', context)

@login_required(login_url='login')
def report_view(request, resolution_id):
    start = time.time()
    r = Resolution.objects.get(pk=resolution_id)
    nome_parte = r.part.part.name
    areas = Area.objects.filter(part=r.part.part)
    report_json = r.statistics
    report_json_dumps = json.dumps(report_json, indent=1, sort_keys=False, ensure_ascii=False)
    
    report = {}
    answers = Answer.objects.filter(resolution=r).order_by("question__section__order")
    done = []
    nr_areas = len(areas) - 3
    nr_total_instrumentos = -3

    report_obj = Report.objects.get(resolution=r)
    report_obj.refresh_report(answers)

    for area in areas.order_by('order'):
        nr_total_instrumentos += area.number_of_instruments
        report[area.name] = {
            'nr_instrumentos': area.number_of_instruments,
        }
        instruments = Instrument.objects.all().order_by('order').filter(area=area)

        for instrument in instruments:
            if instrument.name != None:
                if 'ABVD' in instrument.name:
                    names = ['Atividades Corporais', 'Atividades Motoras', 'Atividades Mentais', 'Atividades Sensoriais']
                    quotations = [report_obj.abvd_atividades_corporais_quotation, report_obj.abvd_atividades_motoras_quotation, report_obj.abvd_atividades_mentais_quotation, report_obj.abvd_atividades_sensoriais_quotation]
                    total = sum(quotations)
                    done.append(instrument.name)
                    print(f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    report[area.name][instrument.name] = {
                        'evaluation': report_obj.abvd_evaluation,
                        'atividades_corporais': report_obj.abvd_atividades_corporais_quotation,
                        'atividades_motoras': report_obj.abvd_atividades_motoras_quotation,
                        'atividades_mentais': report_obj.abvd_atividades_mentais_quotation,
                        'atividades_sensoriais': report_obj.abvd_atividades_sensoriais_quotation,
                        'total': total,
                        'graph': make_graph(names, quotations, 0, instrument.highest_max_quotation),
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'abvd_notes',
                    }
                
                elif 'AIVD' in instrument.name:
                    report[area.name][instrument.name] = {
                        'evaluation': report_obj.aivd_evaluation,
                        'utilizacao_telefone': report_obj.aivd_utilizacao_telefone_quotation,
                        'fazer_compras': report_obj.aivd_fazer_compras_quotation,
                        'preparar_refeicoes': report_obj.aivd_preparacao_refeicoes_quotation,
                        'tarefas_domesticas': report_obj.aivd_tarefas_domesticas_quotation,
                        'lavar_roupa' : report_obj.aivd_lavar_roupa_quotation,
                        'utilizar_transportes': report_obj.aivd_utilizar_transportes_quotation,
                        'manejo_medicacao': report_obj.aivd_manejo_mediacao_quotation,
                        'responsabilidade_financeira': report_obj.aivd_responsabilidades_financeiras_quotation,
                        'respondido': False,
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'aivd_notes',
                    }
                
                elif 'BSI' in instrument.name:
                    names = ['Somatização','Obsessão-Compulsão','Sensibilidade Interpessoal','Depressão','Ansiedade','Hostilidade','Ansiedade Fóbica','Ideação Paranóide','Psicoticismo']
                    quotations = [report_obj.bsi_somatizacao, report_obj.bsi_obssessivo_compulsivo, report_obj.bsi_sensibilidade_interpessoal, report_obj.bsi_depressao, report_obj.bsi_ansiedade, report_obj.bsi_hostilidade, report_obj.bsi_ansiedade_fobica, report_obj.bsi_paranoide, report_obj.bsi_psicoticismo]
                    done.append(instrument.name)
                    print(f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    report[area.name][instrument.name] = {
                        'somatizacao': report_obj.bsi_somatizacao,
                        'obsessao': report_obj.bsi_obssessivo_compulsivo,
                        'depressao': report_obj.bsi_depressao,
                        'sensibilidade_interpessoal': report_obj.bsi_sensibilidade_interpessoal,
                        'ansiedade': report_obj.bsi_ansiedade,
                        'hostilidade': report_obj.bsi_hostilidade,
                        'ansiedade_fobica': report_obj.bsi_ansiedade_fobica,
                        'paranoide': report_obj.bsi_paranoide,
                        'psicoticismo': report_obj.bsi_psicoticismo,
                        'igs': round(report_obj.bsi_igs, 2),
                        'tsp': round(report_obj.bsi_tsp, 2),
                        'isp': round(report_obj.bsi_isp, 2),
                        'graph': make_graph(names, quotations, 0, 28),
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'bsi_notes',
                    }
                
                elif 'ACE-R' in instrument.name:
                    names = ['Atenção e Orientação','Memória','Fluência','Linguagem','Visuo-Espacial']
                    quotations = [report_obj.acer_atencao_orientacao_quotation, report_obj.acer_memoria_quotation, report_obj.acer_fluencia_quotation, report_obj.acer_linguagem_quotation, report_obj.acer_visuo_espacial_quotation]
                    done.append(instrument.name)
                    print(f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    report[area.name][instrument.name] = {
                        'total': sum(quotations),
                        'evaluation': report_obj.acer_evaluation,
                        'acer_atencao_orientacao': report_obj.acer_atencao_orientacao_quotation,
                        'acer_memoria': report_obj.acer_memoria_quotation,
                        'acer_fluencia': report_obj.acer_fluencia_quotation,
                        'acer_linguagem': report_obj.acer_linguagem_quotation,
                        'acer_visuo_espacial': report_obj.acer_visuo_espacial_quotation,
                        'graph': make_graph(names, quotations, 0, instrument.highest_max_quotation),
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'acer_notes',
                    }
                
                elif 'MMSE' in instrument.name:
                    names = ['Atenção e Orientação','Memória','Linguagem','Visuo-Espacial']
                    quotations = [report_obj.mmse_atencao_orientacao_quotation, report_obj.mmse_memoria_quotation, report_obj.mmse_lingua_quotation, report_obj.mmse_visuo_espacial_quotation]
                    total = sum(quotations)
                    done.append(instrument.name)
                    print(f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
                    report[area.name][instrument.name] = {
                        'total': total,
                        'evaluation': report_obj.mmse_evaluation,
                        'mmse_atencao_orientacao': report_obj.mmse_atencao_orientacao_quotation,
                        'mmse_memoria': report_obj.mmse_memoria_quotation,
                        'mmse_linguagem': report_obj.mmse_lingua_quotation,
                        'mmse_visuo_espacial': report_obj.mmse_visuo_espacial_quotation,
                        'graph': make_graph(names, quotations, 0, instrument.highest_max_quotation),
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'mmse_notes',
                    }

                elif 'PANAS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'interessado': report_obj.panas_interessado,
                        'nervoso': report_obj.panas_nervoso,
                        'entusiasmado': report_obj.panas_entusiasmado,
                        'amedrontado': report_obj.panas_amedrontado,
                        'inspirado': report_obj.panas_inspirado,
                        'ativo': report_obj.panas_ativo,
                        'assustado': report_obj.panas_assustado,
                        'culpado': report_obj.panas_culpado,
                        'determinado': report_obj.panas_determinado,
                        'atormentado': report_obj.panas_atormentado,
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'panas_notes',
                    }
                
                elif 'HADS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'ansiedade_evaluation': report_obj.hads_estado_ansiedade_evaluation,
                        'ansiedade_quotation': report_obj.hads_estado_ansiedade_quotation,
                        'depressao_evaluation': report_obj.hads_estado_depressao_evaluation,
                        'depressao_quotation': report_obj.hads_estado_depressao_quotation,
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'hads_notes',
                    }
                
                elif 'GDS' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'gds': report_obj.gds_nivel,
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'gds_notes',
                    }

                elif 'Áreas Complementares' in instrument.name:
                    done.append(instrument.name)
                    report[area.name][instrument.name] = {
                        'memoria_visual_imediata': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('21').get('quotation'),
                        'memoria_visual_diferida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('22').get('quotation'),
                        'atencao_mantida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('23').get('quotation'),
                        'atencao_dividida': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('24').get('quotation'),
                        'orientacao_esq_dir': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('25').get('quotation'),
                        'abstracao_verbal': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('26').get('quotation'),
                        'compreensao_instrucoes': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('27').get('quotation'),
                        'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                        'notes': 'ac_notes',
                    }


                elif 'None' in instrument.name:
                    if area.name == 'Consciência, Humor e Comportamento':
                        cons = MultipleChoicesCheckbox.objects.filter(answer__resolution=r, answer__question__name='Consciência')
                        cons_list = [self.choice.name for self in cons]
                        mot = MultipleChoicesCheckbox.objects.filter(answer__resolution=r, answer__question__name='Atividade Motora')
                        mot_list = [self.choice.name for self in mot]
                        hum = MultipleChoicesCheckbox.objects.filter(answer__resolution=r, answer__question__name='Humor')
                        hum_list = [self.choice.name for self in hum]
                        report[area.name][area.name] = {               
                            'chc_consciencia': ", ".join(cons_list),
                            'chc_motora': ", ".join(mot_list),
                            'chc_humor': ", ".join(hum_list),
                            'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                            'notes': 'chc_notes',
                        }
                    
                    elif area.name == 'Cooperação dada na entrevista':
                        coop = ''
                        respondido = r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {               
                            'respondido': respondido,
                            'notes': 'coop_notes',
                        }
                        if respondido:
                            coop = Answer.objects.filter(resolution=r, question__name='Cooperação dada na entrevista').get().multiple_choice_answer.name
                            report[area.name][area.name]['coop'] = coop
                        
                        
                    
                    elif area.name == 'Relação com o Avaliador':
                        rel = ''
                        respondido = r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0
                        report[area.name][area.name] = {               
                            'respondido': r.statistics.get(f'{area.id}').get(f'{instrument.id}').get('answered') > 0,
                            'notes': 'rel_notes',
                        }
                        if respondido:
                            rel = Answer.objects.filter(resolution=r, question__name='Relação com o Avaliador').get().multiple_choice_answer.name
                            report[area.name][area.name]['rel'] = rel
                        
                        
    
    # for area in areas.order_by('order'):
    #     report[area.id] = {}
    #     instruments = Instrument.objects.all().order_by('order').filter(area=area)
    #     for instrument in instruments:
    #         quotations = []
    #         names = []
    #         report[area.id][instrument.name] = {}
    #         dimensions = Dimension.objects.all().order_by('order').filter(instrument=instrument)
    #         report[area.id][instrument.name]["Total"] = 0
    #         report[area.id][instrument.name]["Graph"] = None
    #         for dimension in dimensions:
    #             if dimension.name != 'None':
    #                 names.append(dimension.name)
    #             report[area.id][instrument.name][dimension.name] = {}
    #             report[area.id][instrument.name][dimension.name]['Total'] = 0
    #             sections = Section.objects.all().order_by('order').filter(dimension=dimension)
    #             for section in sections:
    #                 if dimension.name == 'None' and section.name != 'None':
    #                     names.append(section.name)
    #                 report[area.id][instrument.name][dimension.name][section.name] = {}
    #                 questions = Question.objects.filter(section=section)
    #                 report[area.id][instrument.name][dimension.name][section.name] = \
    #                     report_json[str(area.id)][str(instrument.id)][str(dimension.id)][str(section.id)].get(
    #                         'quotation')
    #                 for question in questions:
    #                     answer = Answer.objects.filter(question=question, resolution=r)
    #                     if answer.exists():
    #                         report[area.id][instrument.name]["Total"] += answer.get().quotation
    #                         report[area.id][instrument.name][dimension.name]["Total"] += answer.get().quotation
    #                         if dimension.name == 'None' and section.name != 'None':
    #                             quotations.append(answer.get().quotation)
    #             if dimension.name != 'None':
    #                 quotations.append(report[area.id][instrument.name][dimension.name]["Total"])

    #         if len(quotations) == len(names) and dimension.name != 'None' or len(quotations) == len(
    #                 names) and section.name != 'None' and len(quotations) != 0:
    #             for answer in answers:
    #                 if answer.instrument == instrument.name and instrument.name not in done:
    #                     print(f"Creating Graph for {instrument.name} ; h-max={instrument.highest_max_quotation}, min={instrument.minimum_quotation}")
    #                     report[area.id][instrument.name]["Graph"] = make_graph(names, quotations,
    #                                                                            0,
    #                                                                            instrument.highest_max_quotation)
    #                     done.append(instrument.name)

    #     for instrument in instruments:
    #         if instrument.name == "BSI":
    #             for answer in answers:
    #                 if answer.instrument == instrument.name:
    #                     names = ['Somatização', 'Obsessões-Compulsões', 'Depressão', 'Sensibilidade Interpessoal',
    #                              'Ansiedade', 'Hostilidade', 'Ansiedade Fóbica', 'Ideação Paranóide', 'Psicotismo']
    #                     quotations = bsi_quotation(answers)
    #                     report[area.id]["BSI"]["Graph"] = make_graph(names, quotations, 0, 28)
    #                     done.append(instrument.name)
    
    #print(json.dumps(r.statistics, indent=1, sort_keys=False, ensure_ascii=False))
    # Funcionalidade
    end = time.time()
    print("Report", (end - start))
    # gera_relatorio_parte()
    context = {'report_json': report_json, 
               'report_json_dumps': report_json_dumps, 
               'report': report, 
               'resolution': r,
               'answers': answers, 
               'instruments': Instrument.objects.all(), 
               'questions': Question.objects.all(),
               'areas': areas, 
               'nome_parte': nome_parte,
               'nr_areas': nr_areas,
               'nr_total_instrumentos': nr_total_instrumentos*2,
               }
    # print(Question.objects.all())
    end = time.time()
    print("Report", (end - start))

    renderizado = render(request, 'protocolo/report2.html', context)
    return renderizado


@login_required(login_url='login')
def report_risk(request):
    start = time.time()
    print(request.POST)
    print("ele chega AQUI")
    lines = []

    risk = Risk.objects.all()
    for r in risk:
        lines.append("Relatório de Risco: ")
        lines.append("Idade:" + str(r.idade))
        lines.append("Sexo:" + str(r.sexo))
        lines.append("Colesterol:" + str(r.colestrol_total))
        lines.append("pressao_arterial:" + str(r.pressao_arterial))
        lines.append("O risco atual é de:" + str(r.risco_de_enfarte))
        lines.append("Comentario:" + str(r.comentario))
    
    context = {'lines': lines
               }
    print("ele chega AQUI2")
    parts_risk = render(request, 'protocolo/parts_risk.html', context)
    print("ele chega AQUI3")
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Risco.pdf"'
    pdf = pisa.pisaDocument(parts_risk.content, response)
    if not pdf.err:
        return response
    
    end = time.time()

    print("Report-risk", (end - start))
    return HttpResponse("Erro ao gerar PDF")

@login_required(login_url='login')
def protocol_participants_view(request, protocol_id):
    doctor = request.user
    protocolo = Protocol.objects.get(pk=protocol_id)
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions, 'protocolo': protocolo}
    return render(request, 'protocolo/protocol-participants.html', context)


@login_required(login_url='login')
def participants_view(request):
    doctor = request.user
    participants = Participante.objects.filter(avaliador=doctor)
    resolutions = Resolution.objects.filter(doctor=doctor)

    context = {'participants': participants, 'resolutions': resolutions}
    return render(request, 'protocolo/participants.html', context)


def login_view(request):
    next = request.GET.get('next')
    if request.method == 'POST':
        next = request.POST.get('next')
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            #print(request.POST)
            next_url = request.POST.get('next')
            if next_url:
                return HttpResponseRedirect(next_url)
            elif request.user.groups.filter(name='Avaliador').exists():
                return redirect('protocolo:dashboard')
            elif request.user.groups.filter(name__in=['Dinamizador','Cuidador','Participante']).exists():
                return redirect('protocolo:diario:dashboard_Care')
                    
    context = {
        'next': next,
    }

    return render(request, 'protocolo/login.html')


from django.contrib.auth import logout

def logout_view(request):
    logout(request)
    #Mudei isto de protocolo/login.html
    return render(request, 'mentha/base.html')


@login_required(login_url='login')
def profile_view(request, participant_id):
    # Falta mostrar as resoluções das partes feitas no perfil e os seus relatorios
    patient = Participante.objects.filter(pk=participant_id).get()
    resolutions = Resolution.objects.filter(patient=patient)
    parteDoUtilizador = ParteDoUtilizador.objects.filter(participante=patient)
    parte = Part.objects.all()
    
    # r = Resolution.objects.filter(patient = patient, part__part__name = "MentHA-Risk")
    existing_risk = ParteDoUtilizador.objects.filter(part__name = "MentHA-Risk")

    print("Existing_risk")
    print(existing_risk)
    print("existing risk")

    user = request.user
    # print(request.user.groups.all())
    
    #isto server para ver qual dos grupos dos users para ver permições e deixar aceder o que ao que
    user_risk = None
    user_tudo = None

    if request.user.groups.filter(name='Administrador').exists():
        user_tudo = request.user.groups.filter(name='Administrador')
    if request.user.groups.filter(name='Avaliador').exists():
        user_tudo = request.user.groups.filter(name='Avaliador' )
    if request.user.groups.filter(name='Avaliador-Risk').exists():
        user_risk = request.user.groups.filter(name='Avaliador-Risk')

    risk_area = Area.objects.get(id = 47)
    pergunta_risk = Question.objects.get(id = 189)
    c = []
    age = calculate_age(patient.nascimento)
    answered_list = []
    percentage_list = []
    
    form = AppointmentForm(request.POST or None)

    if request.method == "POST":
        form = AppointmentForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.participante = patient
            obj.save()
    
    # if request.user.groups.filter(name='Avaliador-Risk').exists():
    #     if not parteDoUtilizador.filter(part= risk_area.part):
    #         parte

    for partesDoUtilizador in parteDoUtilizador:
        resolution = resolutions.filter(part=partesDoUtilizador)
        if not resolution:
            answered_list.append(0)
            percentage_list.append(0)
        else:
            s = resolution.get().statistics
            answered_list.append(s.get('total_answered'))
            percentage_list.append(s.get('total_percentage'))
    for cuidador in Cuidador.objects.all():
        if cuidador in patient.cuidadores.all():
            c.append(cuidador.nome)

    cuidadores = ", ".join(c)
    
    # Copia da patient_overview_view
    areas = Area.objects.order_by('part__order', 'order').all()
    # print(areas)
    overview_list = []

    # Getting all unique area names to display on html page
    for area in areas:
        for instrument in Instrument.objects.filter(area=area):
            text = area.name
            if instrument.name != "None":
                text = text + " - " + instrument.name
            overview_list.append(text)

    ow_l = dict.fromkeys(x for x in overview_list)

    percentages = {}
    # Getting % done per area, per part, to display done or not done on HTML
    for part_do_utilizador in parteDoUtilizador:
        percentages[part_do_utilizador.id]={}
        part = part_do_utilizador.part
        r = resolutions.filter(part=part_do_utilizador)
        if len(r) < 1:
            for text in overview_list:
                this_area = Area.objects.filter(part=part, name=text.split(' - ')[0]).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    percentages[part_do_utilizador.id][text] = "not done"
        else:
            r = r.get()
            for text in overview_list:
                area_text = text.split(' - ')[0]
                this_area = Area.objects.filter(part=part, name=area_text).order_by('order')
                if len(this_area) < 1:
                    percentages[part_do_utilizador.id][text] = "does not exist"
                else:
                    this_area = this_area.get()
                    instruments = Instrument.objects.filter(area=this_area, area__part=part).order_by('order')
                    for instrument in instruments:
                        if instrument.name != 'None':
                            p = r.statistics[f"{this_area.id}"][f"{instrument.id}"].get('percentage')
                            
                        else:
                            p = r.statistics[f"{this_area.id}"].get('percentage')

                        if p == 100:
                            percentages[part_do_utilizador.id][text] = "done"
                        else:
                            percentages[part_do_utilizador.id][text] = "not done"

    context = {'patient': patient, 'cuidadores': cuidadores, 'resolutions': resolutions,'partesdoutilizador': zip(parteDoUtilizador, answered_list, percentage_list),
               'partedoutilizador': parteDoUtilizador,
               'overview_list': ow_l,
               'percentages': percentages, 'age': age,
               'parte': parte,
               'risk_area':risk_area,
               'pergunta_risk':pergunta_risk,
               'form' : form,
               'user' : user,
                'user_risk' : user_risk,
                'user_tudo' : user_tudo,
                'existing_risk': existing_risk,
               }
    return render(request, 'protocolo/profile.html', context)


@login_required(login_url='login')
def gds_overview_view(request, protocol_id, part_id, area_id, instrument_id, patient_id):
    patient = Participante.objects.filter(pk=patient_id).get()
    part = Part.objects.get(pk=part_id)
    r = Resolution.objects.filter(patient=patient, doctor=request.user, part=part).get()
    answers = Answer.objects.filter(resolution=r)
    name_list = ['Respondente #1', 'Respondente #2', 'Respondente #3']
    overview_list = []
    questions = Question.objects.filter(section__dimension__instrument__name="GDS", name__in=name_list
                                        ).order_by('order')

    rowspan_estadio1 = len(questions[0].possible_answers.all().filter(description='1'))
    rowspan_estadio2 = len(questions[0].possible_answers.all().filter(description='2'))
    rowspan_estadio3 = len(questions[0].possible_answers.all().filter(description='3'))
    rowspan_estadio4 = len(questions[0].possible_answers.all().filter(description='4')) + len(
        questions[0].possible_answers.all().filter(description='4 - 2'))
    rowspan_estadio5 = len(questions[0].possible_answers.all().filter(description='5')) + len(
        questions[0].possible_answers.all().filter(description='5 - 2'))
    rowspan_estadio6 = len(questions[0].possible_answers.all().filter(description='6')) + len(
        questions[0].possible_answers.all().filter(description='6 - 2')) + len(
        questions[0].possible_answers.all().filter(description='6 - 3')) + len(
        questions[0].possible_answers.all().filter(description='6 - 4'))
    rowspan_estadio7 = len(questions[0].possible_answers.all().filter(description='7'))

    answers_dict = {'1': {},
                    '2': {},
                    '3': {},
                    }

    for q in questions:
        for pa in q.possible_answers.all():
            respondente = q.name.split('#')[1]
            a = Answer.objects.filter(question=q, resolution=r)
            if len(a) > 0:
                a = a.get()
                if a.MCCAnswer is not None:
                    for mca in a.MCCAnswer.all():
                        if mca.choice == pa:
                            answers_dict[respondente][pa.name] = "checked"

    rowspans = {
        '1': rowspan_estadio1,
        '2': rowspan_estadio2,
        '3': rowspan_estadio3,
        '4': rowspan_estadio4,
        '5': rowspan_estadio5,
        '6': rowspan_estadio6,
        '7': rowspan_estadio7,
    }

    list_2, list_3, list_4, list_5, list_6, list_7 = [], [], [], [], [], []

    for pa in questions[0].possible_answers.all():
        if pa.description == '2':
            list_2.append(pa.name)
        elif pa.description == '3':
            list_3.append(pa.name)
        elif pa.description == '4' or pa.description == "4 - 2":
            list_4.append(pa.name)
        elif pa.description == '5' or pa.description == "5 - 2":
            list_5.append(pa.name)
        elif pa.description == '6' or pa.description == "6 - 2" or pa.description == "6 - 3" or pa.description == "6 - 4":
            list_6.append(pa.name)
        elif pa.description == '7':
            list_7.append(pa.name)

    questions_dict = {'2': list_2,
                      '3': list_3,
                      '4': list_4,
                      '5': list_5,
                      '6': list_6,
                      '7': list_7,
                      }

    max_2, max_3, max_4, max_5, max_6, max_7 = [0, 0, 0], [0, 0, 0], [0, 0, 0], [0, 0, 0], [0, 0, 0], [0, 0, 0]
    for q in questions:
        respondente = q.name.split('#')[1]
        a = Answer.objects.filter(question=q, resolution=r)
        if len(a) > 0:
            a = a.get()
            for mca in a.MCCAnswer.all():
                if mca.choice.description == '2':
                    max_2[int(respondente) - 1] += 1
                elif mca.choice.description == '3':
                    max_3[int(respondente) - 1] += 1
                elif mca.choice.description in ['4', '4 - 2']:
                    max_4[int(respondente) - 1] += 1
                elif mca.choice.description in ['5', '5 - 2']:
                    max_5[int(respondente) - 1] += 1
                elif mca.choice.description in ['6', '6 - 2', '6 - 3', '6 - 4']:
                    max_6[int(respondente) - 1] += 1
                elif mca.choice.description in ['7']:
                    max_7[int(respondente) - 1] += 1
    total_2, total_3, total_4, total_5, total_6, total_7 = 0, 0, 0, 0, 0, 0
    for pa in questions[0].possible_answers.all():
        if pa.description == '2':
            total_2 += 1
        elif pa.description == '3':
            total_3 += 1
        elif pa.description in ['4', '4 - 2']:
            total_4 += 1
        elif pa.description in ['5', '5 - 2']:
            total_5 += 1
        elif pa.description in ['6', '6 - 2', '6 - 3', '6 - 4']:
            total_6 += 1
        elif pa.description in ['7']:
            total_7 += 1

    max_total_dict = {'2': [max(max_2), total_2],
                      '3': [max(max_3), total_3],
                      '4': [max(max_4), total_4],
                      '5': [max(max_5), total_5],
                      '6': [max(max_6), total_6],
                      '7': [max(max_7), total_7],
                      }

    context = {'questions': questions,
               'question': questions[0],
               'rowspans': rowspans,
               'dict': answers_dict,
               'questions_dict': questions_dict,
               'max_total_dict': max_total_dict,
               'protocol_id': protocol_id,
               'part_id': part_id,
               'area_id': area_id,
               'instrument_id': instrument_id,
               'patient_id': patient_id,
               }
    return render(request, 'protocolo/gds_overview.html', context)


@login_required(login_url='login')
def patient_overview_view(request, participant_id):
    patient = Participante.objects.filter(pk=participant_id).get()
    resolutions = Resolution.objects.filter(patient=patient).order_by('part__order')
    areas = Area.objects.order_by('part__order', 'order').all()
    # print(areas)
    parts = Part.objects.all()
    overview_list = []

    # Getting all unique area names to display on html page
    for area in areas:
        for instrument in Instrument.objects.filter(area=area):
            text = area.name
            if instrument.name != "None":
                text = text + " - " + instrument.name
            overview_list.append(text)

    ow_l = dict.fromkeys(x for x in overview_list)

    percentages = {1: {}, 2: {}, 3: {}, 4: {}, 5: {}, 6: {}}
    # Getting % done per area, per part, to display done or not done on HTML
    for part in parts:
        r = resolutions.filter(part=part)
        if len(r) < 1:
            for text in overview_list:
                this_area = Area.objects.filter(part=part, name=text.split(' - ')[0]).order_by('order')
                if len(this_area) < 1:
                    percentages[part.order][text] = "does not exist"
                else:
                    percentages[part.order][text] = "not done"
        else:
            r = r.get()
            for text in overview_list:
                area_text = text.split(' - ')[0]
                this_area = Area.objects.filter(part=part, name=area_text).order_by('order')
                if len(this_area) < 1:
                    percentages[part.order][text] = "does not exist"
                else:
                    this_area = this_area.get()
                    instruments = Instrument.objects.filter(area=this_area, area__part=part).order_by('order')
                    for instrument in instruments:
                        if instrument.name != 'None':
                            p = r.statistics[f"{this_area.id}"][f"{instrument.id}"].get('percentage')
                        else:
                            p = r.statistics[f"{this_area.id}"].get('percentage')

                        if p == 100:
                            percentages[part.order][text] = "done"
                        else:
                            percentages[part.order][text] = "not done"

    # print(percentages.get(2))
    # print(overview_list)
    context = {'patient': patient,
               'resolutions': resolutions,
               'parts': parts,
               'overview_list': ow_l,
               'percentages': percentages,
               }
    return render(request, 'protocolo/patient_overview.html', context)


#quero obter o json com as respostas de um paciente de risk
def risk_json(path,smoking, idade,colesterol,hipertensao):

    print("entrou no risk_json")
    #abrir json risk_men
    data = open_json(path)
    print("entrou no risk_json2")
    for i in data:
        if(i == smoking):
            print("entrou no risk_json3")
            for j in data[i]:
                min=j.split('-')[0]
                max=j.split('-')[1]
                min = int(min)
                max = int(max)
                idade = int(idade)
                if(idade in range(min,max+1)):
                    print("entrou no risk_json4")
                    for k in data[i][j]:
                        min=k.split('-')[0]
                        max=k.split('-')[1]
                        min = int(min)
                        max = int(max)
                        hipertensao = int(hipertensao)
                        if(hipertensao in range(min,max+1)):
                           print("entrou no risk_json5")
                           for l in data[i][j][k]:
                               min=l.split('-')[0]
                               max=l.split('-')[1]
                               min = float(min)
                               max = float(max)
                               print("este é o valor max do colestrol" + str(max))
                               print("este é o valor min do colestrol" + str(min))
                               print("este é o valor do colestrol")
                               print(colesterol)
                               print("este é o valor do colestrol")
                               colesterol = float(colesterol)
                               if(colesterol in float_range(min,max+0.1)):
                                   print("entrou no risk_json6")
                                   print(data[i][j][k][l])
                                   return data[i][j][k][l]

def generate_id():
    # Generate a random number
    random_num = str(random.randint(0, 99999999)).encode()

    # Generate a SHA-256 hash of the random number.
    hash_obj = hashlib.sha256(random_num)
    hex_digit = hash_obj.hexdigest()

    return hex_digit[:10]
#funcao que abre um ficheiro Json e devolve um dicionario
def open_json(file_name):
    with open(file_name) as json_file:
        data = json.load(json_file)
        return data
#funcao que faça um float range de 0.1 em 0.1
def float_range(start, stop, step=0.1):
    while start < stop:
        yield round(start, 1)
        start += step   
def gera_relatorio_risk_pdf(parte_risk,patient, username):
    
    document = Document()
    #conversao para boleans
    if parte_risk.diabetes == "True":
        parte_risk.diabetes = True
    else:
        parte_risk.diabetes = False
    if parte_risk.enfarte == "True":
        parte_risk.enfarte = True
    else:
        parte_risk.enfarte = False
    if parte_risk.doenca_rins == "True":
        parte_risk.doenca_rins = True
    else:
        parte_risk.doenca_rins = False
    if parte_risk.hipercolestrol == "True":
        parte_risk.hipercolestrol = True
    else:
        parte_risk.hipercolestrol = False
    if parte_risk.avc == "True":
        parte_risk.avc = True
    else:
        parte_risk.avc = False
    if parte_risk.doenca_pernas == "True":
        parte_risk.doenca_pernas = True
    else:
        parte_risk.doenca_pernas = False
    #guardar valores para uso
    risco_baixo = ""
    hipercolesterol = ""
    drc = ""
    fuma = ""
    cardio = ""
    diabetes = ""

    #conversao para ints
    idade = int(parte_risk.idade)
    risco = int(parte_risk.risco_de_enfarte)
    #cores
    color_low = '00FF00'  # Green
    color_moderate = 'FFD700'  # Yellow 'FFFF00'
    color_high = 'FF0000'  # Red
    color_very_high = '8B0000'  # Dark Red

    # Cabeçalho
    paragraph = document.add_paragraph(f'MentHA-Risk')


    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(f'Avaliação de risco Cardiovascular de {patient.__str__()}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Relatório


    #Dados Sócio-demográficos 
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados Socio-demográficos')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Nome:')
    run.bold = True
    paragraph.add_run(f' {patient.__str__()}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Instituição de Referência:')
    run.bold = True
    paragraph.add_run(f' {patient.referenciacao}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Local de Residência:')
    run.bold = True
    paragraph.add_run(f' {patient.localizacao}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'E-mail:')
    run.bold = True
    paragraph.add_run(f' {patient.email}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    #Dados para SCORE-2
    paragraph = document.add_paragraph()  
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados para SCORE-2')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Idade:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.idade}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Sexo:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.sexo}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if parte_risk.fumador == 'smoking':
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'É fumador: ')
        run.bold = True
        paragraph.add_run(f'Sim')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        fuma = "É fumador"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'É fumador:')
        run.bold = True
        paragraph.add_run(f' Não')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
        fuma = "Não fuma"

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Pressão arterial sistólica:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.pressao_arterial} mmHg')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol Não HDL:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_nao_hdl} mmmol/L')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY



    #Dados Complementares
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run('Dados Complementares')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Peso:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.peso} kg')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'IMC:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.imc}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Altura:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.altura} cm')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol total:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_total} mmol/L')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Colesterol HDL:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.colestrol_hdl} mmol/L')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Hemoglobina_gliciada:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.hemoglobina_gliciada} %')  
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if parte_risk.diabetes:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Diabetes:')
        run.bold = True
        paragraph.add_run(f' Tem Diabetes')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        diabetes = "com Diabetes"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Diabetes:')
        run.bold = True
        paragraph.add_run(f' Não tem Diabetes')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        diabetes = "sem Diabetes"
    if parte_risk.diabetes:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Há quantos anos o paciente tem diabétes:')
        run.bold = True
        paragraph.add_run(f' {parte_risk.anos_diabetes} anos')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        pass

    if parte_risk.avc:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum AVC (Acidente Vascular Cerebral):')
        run.bold = True
        paragraph.add_run(f' Sim teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum AVC (Acidente Vascular Cerebral):')
        run.bold = True
        paragraph.add_run(f' Não teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.enfarte:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum Enfarte (Ataque cardiaco):')
        run.bold = True
        paragraph.add_run(f' Sim teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        cardio = "com histório de evento cardiovascular major"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Teve algum Enfarte (Ataque cardiaco):')
        run.bold = True
        paragraph.add_run(f' Não teve')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        cardio = "sem histório de evento cardiovascular major"
    if parte_risk.doenca_rins:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença dos Rins:')
        run.bold = True
        paragraph.add_run(f' Sim tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        drc = "e com doença renal crónica"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença dos Rins:')
        run.bold = True
        paragraph.add_run(f' Não tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        drc = "sem doença renal crónica"
    if parte_risk.doenca_pernas:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença das Pernas:')
        run.bold = True
        paragraph.add_run(f' Sim tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Tem doença das Pernas:')
        run.bold = True
        paragraph.add_run(f' Não tem')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if parte_risk.hipercolestrol:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Historial de hipercolesterolemia familiar:')
        run.bold = True
        paragraph = document.add_paragraph(f' Existe familiares com hipercolesterolémia')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        hipercolesterol = "com hipercolesterolémia familiar"
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Historial de hipercolesterolemia familiar:')
        run.bold = True
        paragraph = document.add_paragraph(f' Não existe familiares com hipercolesterolémia')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        hipercolesterol = "sem hipercolesterolémia familiar"
    #Resumo dos dados/Resultados
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Resultados')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Probabilidade de ter um Risco Cardiovascular:')
    run.bold = True
    paragraph.add_run(f' {parte_risk.risco_de_enfarte} %')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if ((risco <= 2 and idade < 50) or (risco < 5 and (idade >= 50 and idade <= 69)) or (risco <= 5 and idade >= 70)):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
        run.bold = True
        paragraph.add_run(f' Baixo')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        add_cores(paragraph, color_low)
        risco_baixo = "Baixo"
    elif (((risco > 2 or risco <= 8) and idade < 50) or (risco < 5 and (idade >= 50 or idade <= 69)) or (risco <= 5 and idade >= 70) ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
        run.bold = True
        paragraph.add_run(f' Moderado')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        add_cores(paragraph, color_moderate)
        risco_baixo = "Moderado"
    elif (((risco >= 2 and risco < 8) and idade<50) or ((risco >= 5 and risco < 10) and (idade >= 50 and idade <= 69)) or((risco >= 7 and risco < 15) and idade >= 70)):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
        run.bold = True
        paragraph.add_run(f' Alto')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        add_cores(paragraph, color_high)
        risco_baixo = "Alto"
    elif ((risco >= 7 and idade < 50) or (risco >= 10 and (idade >= 50 and idade <=69)) or(risco >= 15 and idade >= 70)):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'Classificação do Risco Cardiovascular:')
        run.bold = True
        paragraph.add_run(f' Muito Alto')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        add_cores(paragraph, color_very_high)
        risco_baixo = "Elevado"

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Descrição: ')
    run.bold = True
    if parte_risk.sexo == 'M':
        paragraph.add_run(f'Um Homen de {idade} anos {cardio},{diabetes} ,{hipercolesterol} e {drc}. O valor de colesterol total de {parte_risk.colestrol_total} mg/dl e o colesterol HDL de {parte_risk.colestrol_hdl} mg/dl; logo o colesterol não HDL de {parte_risk.colestrol_nao_hdl} mg/dl. {fuma} e o valor da pressão arterial sistolica é de {parte_risk.pressao_arterial} mmHg. O score de risco é de {parte_risk.risco_de_enfarte}% o que significa que tem {risco_baixo} risco cardiovascular.')
    else:
        paragraph.add_run(f'Uma Mulher de {idade} anos {cardio},{diabetes} ,{hipercolesterol} e {drc}. O valor de colesterol total de {parte_risk.colestrol_total} mg/dl e o colesterol HDL de {parte_risk.colestrol_hdl} mg/dl; logo o colesterol não HDL de {parte_risk.colestrol_nao_hdl} mg/dl. {fuma} e o valor da pressão arterial sistolica é de {parte_risk.pressao_arterial} mmHg. O score de risco é de {parte_risk.risco_de_enfarte}% o que significa que tem {risco_baixo} risco cardiovascular.')
    
    if idade <70 and idade >= 40:

        # nome_ficheiro_imagem = 'SCORE-2-1-' +patient.__str__()+generate_id()+'.png'
        img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\SCORE-2-1.png')
        # new_img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\img-report\SCORE-2-1-'+patient.__str__()+generate_id()+'.png')

        document.add_picture(img_path, width=Inches(4.5), height=Inches(4.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        

    elif idade >= 70 and idade < 90:

        # nome_ficheiro_imagem = 'SCORE-2-1-' +patient.__str__()+generate_id()+'.png'
        img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\SCORE-2-90.png')
        
        document.add_picture(img_path, width=Inches(4.5), height=Inches(4.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #fazer parte dinamica para mostrar os valores de risco

    #Recomendações
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Recomendações')
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'O Paciente deve:')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Deixar de fumar')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de álcool')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de sal')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de gorduras saturadas e colesterol')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Reduzir o consumo de açúcares e doces')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Fazer uma alimentação saudável e equilibrada')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'  •	Praticar exercício físico regularmente')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #Comentário do Avaliador
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Comentário do Avaliador:')
    run.bold = True
    run.font.size = Pt(14)
    paragraph = document.add_paragraph(f'{parte_risk.comentario}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add the image
    # img_path = os.path.join(os.getcwd(), 'img', 'example.jpg')
    # document.add_picture(img_path, width=Inches(3), height=Inches(3))asdsad
    
    print("CHEGA A  ZIMBORA")
    # Assinatura
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph(f'O avaliador, {username}')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph = document.add_paragraph()

    #imagens logos no fundo do documento
    # cwd = os.getcwd()
    # cwd2 = os.path.join(cwd, 'mentha', 'static', 'img', 'img-logo','ulht.png')
    section = document.sections[0]

    # Configurar o rodapé
    footer = section.footer
    footer.is_linked_to_previous = False  # Certifique-se de que o rodapé não esteja vinculado ao anterior

# Criar parágrafo vazio no rodapé para adicionar as imagens
    paragraph = footer.paragraphs[0]
    image_paths = [
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'ulht.png'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'dgs_footer.png'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'adebe.png'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'copelabs.jpg'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'gira.png'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'cvp.jpg'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'familiarmente.jpg'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'elo.png'),
    os.path.join(os.getcwd(), 'protocolo', 'static', 'protocolo', 'img', 'img-logo', 'mentha-logo.png')
    ]
    for i,image_path in enumerate(image_paths):
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(0.3), height=Inches(0.3))
        if i < len(image_path)-1:
            run.add_text(" ")

    # img_path6 = os.path.join(os.getcwd(), 'mentha\\static\mentha\\pareceiros_sm\\dgs_footer.png')
    # document.add_picture(img_path6, width=Inches(1.5), height=Inches(1.5))
    # img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\logo4.png')
    # document.add_picture(img_path, width=Inches(1.5), height=Inches(1.5))
    # img_path = os.path.join(os.getcwd(), 'protocolo\static\protocolo\img\logo5.png')
    # document.add_picture(img_path, width=Inches(1.5), height=Inches(1.5))
    
    # Save the Word document
    nome_ficheiro = 'Risco_Cardiovascular' +'_'+ patient.__str__() + generate_id() 
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    print(docx_path)
    document.save(docx_path)
    print("CHEGA A  ZIMBORA2")
    # Convert the Word document to PDF

    pdf_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.pdf')
    pythoncom.CoInitialize()
    
    
    print("CHEGA A  ZIMBORA3")
    word_app = win32.gencache.EnsureDispatch('Word.Application')
    doc = word_app.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)
    doc.Close()
    word_app.Quit()
    print("CHEGA A  ZIMBORA4")
    # Create a Django File object from the PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = io.BytesIO(f.read())

    # Create a Django File object from the Word document
    with open(docx_path, 'rb') as f:
        docx_data = io.BytesIO(f.read())
        
    # Assign the PDF file to the file field of sessaoDoGrupo
    parte_risk.relatorio.save(f'{nome_ficheiro}.pdf', pdf_data)
    parte_risk.save()
    parte_risk.relatorio_word.save(f'{nome_ficheiro}.docx', docx_data)
    parte_risk.save()
    print("CHEGA A  ZIMBORA5")
    # Delete the temporary files
    os.remove(docx_path)
    os.remove(pdf_path)


def gera_relatorio_parte(resolution, chc, coop, rel):
    # chc vai ser uma lista de listas

    document = Document()
    # Cabeçalho
    paragraph = document.add_paragraph(f'Relatório do Protocolo de Avaliação MentHA')

    # para pôr em itálico (chato... talvez exista algo melhor)
    for run in paragraph.runs:
        run.font.italic = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = document.add_heading(f'{resolution.part.part.name} de {resolution.patient.info_sensivel.nome}', 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabela
    nr_areas = resolution.part.part.number_of_areas
    areas = resolution.part.part.area
    table = document.add_table(rows=1, cols=3, style="Table Grid")

    # Cabeçalho da Tabela
    heading_row = table.rows[0].cells
    heading_row[0].text = "Área de Avaliação"
    heading_row[1].text = "Resultado"

    for area in areas:
        name = area.name
        # add new row to table
        data_row = table.add_row().cells

        # add headings
        if name == "Questionário Sociodemográfico":
            continue

        data_row[0].text = area.name

        if name == 'Consciência, Humor e Comportamento':
            chc_consciencia = ", ".join(chc[0])
            chc_motora = ", ".join(chc[1])
            chc_humor = ", ".join(chc[2])

            if len(chc_consciencia) > 0:
                p = data_row[1].paragraphs[0]
                p.add_run('Consciência: ').bold = True
                p.add_run(chc_consciencia)

            if len(chc_motora) > 0:
                p = data_row[1].add_paragraph()
                p.add_run('Atividade Motora: ').bold = True
                p.add_run(chc_motora)

            if len(chc_humor) > 0:
                p = data_row[1].add_paragraph()
                p.add_run('Humor: ').bold = True
                p.add_run(chc_humor)
        
        if name == 'Cooperação dada na Entrevista':
             if len(coop) > 0:
                p = data_row[1].paragraphs[0]
                p.add_run(coop)

        if name == 'Relação com o Avaliador':
             if len(rel) > 0:
                p = data_row[1].paragraphs[0]
                p.add_run(rel)
                
        #data_row[1].text = "Texto"

    # Assinatura
    paragraph = document.add_paragraph(f'O avaliador, {resolution.doctor.username}')

    # Save the Word document
    nome_ficheiro = 'Risco_Cardiovascular' +'_'+ resolution.patient.info_sensivel.nome
    nome_ficheiro = nome_ficheiro.replace(" ", "")
    docx_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.docx')
    print(docx_path)
    document.save(docx_path)
    # Convert the Word document to PDF

    pdf_path = os.path.join(os.getcwd(), f'{nome_ficheiro}.pdf')
    pythoncom.CoInitialize()
    
    
    word_app = win32.gencache.EnsureDispatch('Word.Application')
    doc = word_app.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)
    doc.Close()
    word_app.Quit()
    # Create a Django File object from the PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = io.BytesIO(f.read())

    # Delete the temporary files
    # os.remove(docx_path)
    # os.remove(pdf_path)


def word ():
    return 1


#funcao para calcular o IMC
def calcular_imc(peso,altura):
    #altura em metros
    #peso em kg
    peso = int(peso)
    altura = int(altura)
    imc = peso/(altura*altura)
    return imc
#funcao para por cores no word
def add_cores(paragraph, color):
    run = paragraph.add_run()
    run.text = u'\u25CF'  # Circle character
    run.font.color.rgb = RGBColor.from_string(color)